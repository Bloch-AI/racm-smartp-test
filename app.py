from flask import Flask, render_template, request, jsonify, send_from_directory, send_file, session, redirect, url_for, g
from dotenv import load_dotenv
from werkzeug.utils import secure_filename
import json
import logging
import os
import re
import uuid
import threading
import anthropic
import mimetypes
from datetime import datetime

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Development mode flag - set DEV_MODE=true for development features
DEV_MODE = os.environ.get('DEV_MODE', 'false').lower() in ('true', '1', 'yes')

# Constants
CLAUDE_MODEL = "claude-sonnet-4-20250514"
MAX_CHAT_HISTORY = 20  # Maximum messages to retain in chat history

# Text extraction imports
try:
    import pdfplumber
    PDF_SUPPORT = True
except ImportError:
    PDF_SUPPORT = False

try:
    from docx import Document as DocxDocument
    DOCX_SUPPORT = True
except ImportError:
    DOCX_SUPPORT = False

try:
    from openpyxl import load_workbook
    XLSX_SUPPORT = True
except ImportError:
    XLSX_SUPPORT = False

load_dotenv()
from database import get_db, RACMDatabase
from auth import (
    get_current_user, require_login, require_admin, require_audit_access,
    login_user, logout_user, check_password, hash_password,
    get_user_accessible_audits, get_active_audit_id, set_active_audit,
    get_user_role, can_edit_record, can_transition_record, get_record_permissions,
    require_record_edit, require_transition, require_non_viewer
)


# ==================== TEXT EXTRACTION ====================

def extract_text_from_file(filepath: str, mime_type: str = None) -> str:
    """Extract text content from various file types.

    Supported: PDF, DOCX, XLSX, TXT, CSV
    Returns extracted text or empty string if extraction fails.
    """
    if not os.path.exists(filepath):
        return ""

    # Determine file type from extension if mime_type not provided
    ext = filepath.rsplit('.', 1)[-1].lower() if '.' in filepath else ''

    try:
        # Plain text files
        if ext in ('txt', 'csv', 'eml') or (mime_type and 'text' in mime_type):
            with open(filepath, 'r', encoding='utf-8', errors='ignore') as f:
                return f.read()[:50000]  # Limit to 50KB of text

        # PDF files
        if ext == 'pdf' or (mime_type and 'pdf' in mime_type):
            if not PDF_SUPPORT:
                return "[PDF extraction not available - pdfplumber not installed]"
            text_parts = []
            with pdfplumber.open(filepath) as pdf:
                for i, page in enumerate(pdf.pages[:50]):  # Limit to 50 pages
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(f"[Page {i+1}]\n{page_text}")
            return "\n\n".join(text_parts)[:50000]

        # Word documents
        if ext == 'docx' or (mime_type and 'wordprocessingml' in mime_type):
            if not DOCX_SUPPORT:
                return "[DOCX extraction not available - python-docx not installed]"
            doc = DocxDocument(filepath)
            text_parts = []
            for para in doc.paragraphs[:500]:  # Limit paragraphs
                if para.text.strip():
                    text_parts.append(para.text)
            # Also extract from tables
            for table in doc.tables[:20]:
                for row in table.rows:
                    row_text = ' | '.join(cell.text.strip() for cell in row.cells if cell.text.strip())
                    if row_text:
                        text_parts.append(row_text)
            return "\n".join(text_parts)[:50000]

        # Excel files
        if ext == 'xlsx' or (mime_type and 'spreadsheetml' in mime_type):
            if not XLSX_SUPPORT:
                return "[XLSX extraction not available - openpyxl not installed]"
            wb = load_workbook(filepath, read_only=True, data_only=True)
            text_parts = []
            for sheet_name in wb.sheetnames[:10]:  # Limit sheets
                sheet = wb[sheet_name]
                text_parts.append(f"[Sheet: {sheet_name}]")
                for row_num, row in enumerate(sheet.iter_rows(max_row=200, values_only=True)):
                    if row_num > 200:
                        break
                    row_values = [str(cell) if cell is not None else '' for cell in row]
                    if any(v.strip() for v in row_values):
                        text_parts.append(' | '.join(row_values))
            wb.close()
            return "\n".join(text_parts)[:50000]

        # Old Word format (.doc)
        if ext == 'doc':
            return "[.doc format not supported - please convert to .docx]"

        # Old Excel format (.xls)
        if ext == 'xls':
            return "[.xls format not supported - please convert to .xlsx]"

        # MSG email files
        if ext == 'msg':
            return "[.msg email extraction not yet implemented]"

        # Images - no text extraction
        if ext in ('png', 'jpg', 'jpeg', 'gif') or (mime_type and 'image' in mime_type):
            return "[Image file - no text content]"

        # ZIP files
        if ext == 'zip':
            return "[ZIP archive - cannot extract text from compressed files]"

        return f"[Unsupported file type: {ext}]"

    except Exception as e:
        return f"[Error extracting text: {str(e)}]"

app = Flask(__name__)

# Secret key configuration with production safety check
app.secret_key = os.environ.get('SECRET_KEY')
if not app.secret_key:
    if DEV_MODE:
        app.secret_key = 'smartpapers-dev-key-DO-NOT-USE-IN-PRODUCTION'
        logger.warning("Using development secret key - set SECRET_KEY env var for production")
    else:
        raise ValueError("SECRET_KEY environment variable is required in production")

# Session security configuration
app.config['SESSION_COOKIE_SECURE'] = os.environ.get('SESSION_COOKIE_SECURE', 'False').lower() == 'true'
app.config['SESSION_COOKIE_HTTPONLY'] = True
app.config['SESSION_COOKIE_SAMESITE'] = 'Lax'

# Claude API key - set via environment variable
ANTHROPIC_API_KEY = os.environ.get('ANTHROPIC_API_KEY', '')

# Rate limiting for login attempts (in-memory, simple implementation)
LOGIN_ATTEMPTS = {}  # {ip: {'count': int, 'last_attempt': datetime}}
MAX_LOGIN_ATTEMPTS = 5
LOGIN_LOCKOUT_MINUTES = 15

# File upload configuration
UPLOAD_FOLDER = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'uploads')
ALLOWED_EXTENSIONS = {'pdf', 'doc', 'docx', 'xls', 'xlsx', 'msg', 'eml', 'txt', 'csv', 'png', 'jpg', 'jpeg', 'gif', 'zip'}
MAX_CONTENT_LENGTH = 50 * 1024 * 1024  # 50MB max file size

app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
app.config['MAX_CONTENT_LENGTH'] = MAX_CONTENT_LENGTH

# Ensure upload folder exists
os.makedirs(UPLOAD_FOLDER, exist_ok=True)


# ==================== Security Middleware ====================

@app.after_request
def add_security_headers(response):
    """Add security headers to all responses."""
    # Prevent MIME type sniffing
    response.headers['X-Content-Type-Options'] = 'nosniff'
    # Prevent clickjacking
    response.headers['X-Frame-Options'] = 'SAMEORIGIN'
    # Enable XSS filter
    response.headers['X-XSS-Protection'] = '1; mode=block'
    # Referrer policy
    response.headers['Referrer-Policy'] = 'strict-origin-when-cross-origin'
    # Permissions policy (disable sensitive features)
    response.headers['Permissions-Policy'] = 'geolocation=(), microphone=(), camera=()'
    return response


def check_rate_limit(ip_address: str) -> tuple:
    """Check if IP is rate limited for login attempts.
    Returns (is_allowed, remaining_attempts, lockout_seconds)
    """
    from datetime import timedelta
    now = datetime.now()

    if ip_address in LOGIN_ATTEMPTS:
        attempt_data = LOGIN_ATTEMPTS[ip_address]
        lockout_end = attempt_data['last_attempt'] + timedelta(minutes=LOGIN_LOCKOUT_MINUTES)

        if attempt_data['count'] >= MAX_LOGIN_ATTEMPTS:
            if now < lockout_end:
                remaining_seconds = (lockout_end - now).total_seconds()
                return False, 0, int(remaining_seconds)
            else:
                # Lockout expired, reset
                LOGIN_ATTEMPTS[ip_address] = {'count': 0, 'last_attempt': now}

    return True, MAX_LOGIN_ATTEMPTS - LOGIN_ATTEMPTS.get(ip_address, {}).get('count', 0), 0


def record_login_attempt(ip_address: str, success: bool):
    """Record a login attempt for rate limiting."""
    now = datetime.now()

    if success:
        # Clear attempts on successful login
        LOGIN_ATTEMPTS.pop(ip_address, None)
    else:
        if ip_address not in LOGIN_ATTEMPTS:
            LOGIN_ATTEMPTS[ip_address] = {'count': 0, 'last_attempt': now}
        LOGIN_ATTEMPTS[ip_address]['count'] += 1
        LOGIN_ATTEMPTS[ip_address]['last_attempt'] = now


@app.context_processor
def inject_audit_context():
    """Inject audit-related variables into all templates."""
    return get_audit_context(auto_select=False)

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

# Initialize database
db = get_db()


# ==================== Error Response Helpers ====================

def error_response(message: str, status_code: int = 400):
    """Return a standardized JSON error response."""
    return jsonify({'error': message}), status_code


def not_found_response(entity: str = 'Resource'):
    """Return a standardized 404 response."""
    return jsonify({'error': f'{entity} not found'}), 404


def get_audit_context(auto_select: bool = True) -> dict:
    """Get the current audit context for the user.

    Returns dict with:
        - accessible_audits: List of audits the user can access
        - active_audit_id: ID of the currently selected audit
        - active_audit: Full details of the active audit (or None)

    If auto_select is True and no audit is selected, selects the first available.
    """
    accessible_audits = get_user_accessible_audits()
    active_audit_id = get_active_audit_id()

    # Auto-select first audit if none selected
    if auto_select and accessible_audits and not active_audit_id:
        set_active_audit(accessible_audits[0]['id'])
        active_audit_id = accessible_audits[0]['id']

    # Find active audit details
    active_audit = None
    if active_audit_id and accessible_audits:
        for audit in accessible_audits:
            if audit['id'] == active_audit_id:
                active_audit = audit
                break

    return {
        'accessible_audits': accessible_audits,
        'active_audit_id': active_audit_id,
        'active_audit': active_audit
    }


def check_attachment_permission(record: dict, record_type: str = 'record') -> tuple:
    """Check if current user can upload attachments to a record.

    Returns (allowed: bool, error_response: tuple or None).
    If allowed is True, error_response is None.
    If allowed is False, error_response is a (jsonify, status_code) tuple.
    """
    user = get_current_user()
    if user.get('is_admin'):
        return True, None

    audit_id = record.get('audit_id') or get_active_audit_id()
    audit = db.get_audit(audit_id) if audit_id else None
    audit_context = {
        'id': audit_id,
        'auditor_id': audit.get('auditor_id') if audit else None,
        'reviewer_id': audit.get('reviewer_id') if audit else None
    }
    if not can_edit_record(user, audit_context, record):
        return False, (jsonify({'error': f'Cannot upload attachment - {record_type} is not editable'}), 403)
    return True, None


def attachment_upload_response(file_info: dict, attachment_id: int) -> dict:
    """Format a standard attachment upload response."""
    return {
        'status': 'uploaded',
        'id': attachment_id,
        'filename': file_info['original_filename'],
        'size': file_info['file_size'],
        'text_extracted': len(file_info['extracted_text']) > 0 and not file_info['extracted_text'].startswith('[')
    }


# Seed initial data if database is empty
def seed_initial_data():
    """Add sample data if database is empty."""
    if not db.get_all_risks():
        db.create_risk('R001', 'Unauthorized system access', 'C001', 'IT Security', '', '', '', '', 'Not Complete', 0, '', 0, 0)
        db.create_risk('R002', 'Data integrity issues', 'C002', 'Data Team', '', '', '', '', 'Effective', 0, '', 0, 0)
        db.create_risk('R003', 'Segregation of duties', 'C003', 'HR/IT', '', '', '', '', 'Not Effective', 1, 'John', 1, 0)

    if not db.get_all_tasks():
        db.create_task('Define audit scope', 'Identify key processes and risks', 'medium', '', 'planning')
        db.create_task('Request documentation', 'PBC list to client', 'medium', '', 'planning')
        db.create_task('Control walkthroughs', 'Document process flows', 'medium', '', 'fieldwork')

# Only seed sample data in development mode
if DEV_MODE:
    seed_initial_data()
    logger.info("DEV_MODE: Seeded initial sample data")

# Data version for frontend refresh detection (thread-safe)
_data_version = 0
_data_version_lock = threading.Lock()


def get_data_version() -> int:
    """Get current data version (thread-safe)."""
    with _data_version_lock:
        return _data_version


def increment_data_version() -> int:
    """Increment and return new data version (thread-safe)."""
    global _data_version
    with _data_version_lock:
        _data_version += 1
        return _data_version


# ==================== AUDIT LIBRARY PROCESSING ====================

LIBRARY_FOLDER = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'library')
os.makedirs(LIBRARY_FOLDER, exist_ok=True)

# Embedding model (lazy loaded)
_embedding_model = None

def get_embedding_model():
    """Lazy load the sentence transformer model."""
    global _embedding_model
    if _embedding_model is None:
        try:
            from sentence_transformers import SentenceTransformer
            _embedding_model = SentenceTransformer('all-MiniLM-L6-v2')
            logger.info("[Library] Loaded embedding model: all-MiniLM-L6-v2")
        except Exception as e:
            logger.warning(f"[Library] Could not load embedding model: {e}")
    return _embedding_model

def generate_embedding(text: str) -> list:
    """Generate embedding for text using sentence transformer."""
    model = get_embedding_model()
    if model is None:
        return None
    try:
        embedding = model.encode(text, convert_to_numpy=True)
        return embedding.tolist()
    except Exception as e:
        logger.error(f"[Library] Error generating embedding: {e}")
        return None

def chunk_document(text: str, chunk_size: int = 500, overlap: int = 50) -> list:
    """Split document text into overlapping chunks.

    Args:
        text: Full document text
        chunk_size: Target words per chunk
        overlap: Words to overlap between chunks

    Returns:
        List of dicts with 'content', 'section', 'token_count'
    """
    chunks = []

    # Split into paragraphs first
    paragraphs = [p.strip() for p in text.split('\n\n') if p.strip()]

    current_chunk = []
    current_words = 0
    current_section = ""

    for para in paragraphs:
        # Try to detect section headers (lines that look like headers)
        lines = para.split('\n')
        for line in lines:
            # Check if this looks like a header (short, possibly numbered)
            if len(line) < 100 and (
                line.isupper() or
                line.startswith(('Chapter', 'Section', 'CHAPTER', 'SECTION')) or
                (len(line.split()) < 8 and line.endswith(':')) or
                (len(line) > 0 and line[0].isdigit() and '.' in line[:5])
            ):
                current_section = line.strip()

        words = para.split()
        word_count = len(words)

        # If adding this paragraph exceeds chunk size, save current chunk
        if current_words + word_count > chunk_size and current_chunk:
            chunk_text = '\n\n'.join(current_chunk)
            chunks.append({
                'content': chunk_text,
                'section': current_section,
                'token_count': current_words
            })

            # Keep overlap from end of previous chunk
            overlap_text = ' '.join(chunk_text.split()[-overlap:]) if overlap > 0 else ''
            current_chunk = [overlap_text] if overlap_text else []
            current_words = len(overlap_text.split())

        current_chunk.append(para)
        current_words += word_count

    # Don't forget the last chunk
    if current_chunk:
        chunk_text = '\n\n'.join(current_chunk)
        chunks.append({
            'content': chunk_text,
            'section': current_section,
            'token_count': len(chunk_text.split())
        })

    return chunks

def process_library_document(filepath: str, doc_id: int, mime_type: str = None) -> int:
    """Process uploaded library document: extract text, chunk, embed, store.

    Returns number of chunks created.
    """
    # Extract text
    text = extract_text_from_file(filepath, mime_type)
    # Check for error messages (e.g., "[PDF extraction not available...]")
    # but allow normal page markers like "[Page 1]"
    if not text or (text.startswith('[') and 'not available' in text.lower()):
        logger.warning(f"[Library] Could not extract text from document {doc_id}")
        return 0

    # Normalize whitespace while preserving paragraph breaks
    # Replace multiple spaces/tabs with single space, but keep newlines for chunking
    import re
    text = re.sub(r'[ \t]+', ' ', text)  # Multiple spaces/tabs -> single space
    text = re.sub(r'\n{3,}', '\n\n', text)  # 3+ newlines -> double newline

    # Chunk the document
    chunks = chunk_document(text)
    logger.info(f"[Library] Document {doc_id}: Created {len(chunks)} chunks")

    # Initialize vector table
    db._init_vector_table()

    # Store chunks with embeddings
    for i, chunk in enumerate(chunks):
        embedding = generate_embedding(chunk['content'])
        db.add_library_chunk(
            document_id=doc_id,
            chunk_index=i,
            content=chunk['content'],
            section=chunk['section'],
            token_count=chunk['token_count'],
            embedding=embedding
        )

    # Update document with chunk count
    db.update_library_document(doc_id, total_chunks=len(chunks))

    return len(chunks)


# ==================== AUTHENTICATION ====================

@app.route('/login', methods=['GET', 'POST'])
def login():
    """Handle user login."""
    error = request.args.get('error')

    if request.method == 'POST':
        # Rate limiting check
        client_ip = request.remote_addr or '127.0.0.1'
        is_allowed, remaining, lockout_seconds = check_rate_limit(client_ip)

        if not is_allowed:
            minutes = lockout_seconds // 60
            return render_template('login.html',
                error=f'Too many failed attempts. Please try again in {minutes} minutes.')

        email = request.form.get('email', '').strip().lower()
        password = request.form.get('password', '')

        if not email or not password:
            return render_template('login.html', error='Please enter email and password')

        user = db.get_user_by_email(email)

        if not user:
            record_login_attempt(client_ip, success=False)
            return render_template('login.html', error='Invalid email or password')

        if not check_password(password, user['password_hash']):
            record_login_attempt(client_ip, success=False)
            return render_template('login.html', error='Invalid email or password')

        if not user['is_active']:
            return render_template('login.html', error='Your account has been deactivated')

        # Login successful - clear rate limit
        record_login_attempt(client_ip, success=True)
        login_user(user['id'], user['email'], user['name'], user['is_admin'])

        # Set default active audit if user has access to any
        accessible_audits = get_user_accessible_audits()
        if accessible_audits and not get_active_audit_id():
            set_active_audit(accessible_audits[0]['id'])

        # Redirect to original destination or home (prevent open redirect)
        next_url = request.args.get('next', '')
        if next_url and not next_url.startswith('/'):
            next_url = url_for('index')  # Reject absolute URLs
        next_url = next_url or url_for('index')
        return redirect(next_url)

    return render_template('login.html', error=error)


@app.route('/logout')
def logout():
    """Log out the current user."""
    logout_user()
    return redirect(url_for('login'))


@app.route('/api/auth/me', methods=['GET'])
def api_auth_me():
    """Get current user info."""
    user = get_current_user()
    if not user:
        return jsonify({'authenticated': False}), 401
    return jsonify({
        'authenticated': True,
        'user': {
            'id': user['id'],
            'email': user['email'],
            'name': user['name'],
            'is_admin': user['is_admin']
        }
    })


@app.route('/api/auth/set-audit', methods=['POST'])
@require_login
def api_set_active_audit():
    """Set the active audit for the current session."""
    data = request.get_json()
    audit_id = data.get('audit_id')

    if not audit_id:
        return jsonify({'error': 'audit_id required'}), 400

    if set_active_audit(audit_id):
        return jsonify({'status': 'ok', 'active_audit_id': audit_id})
    else:
        return jsonify({'error': 'Access denied to this audit'}), 403


@app.route('/api/auth/accessible-audits', methods=['GET'])
@require_login
def api_accessible_audits():
    """Get list of audits accessible to current user."""
    audits = get_user_accessible_audits()
    return jsonify({
        'audits': audits,
        'active_audit_id': get_active_audit_id()
    })


# ==================== PAGE ROUTES ====================

@app.route('/')
@require_login
def index():
    ctx = get_audit_context()

    # Get user role for template (viewers don't see AI link)
    user = get_current_user()
    user_role = get_user_role(user) if user else None

    return render_template('index.html',
                           active_page='workpapers',
                           user_role=user_role,
                           **ctx)

# ==================== RACM (Spreadsheet) API ====================

@app.route('/api/data', methods=['GET'])
@require_login
def get_data():
    """Get all spreadsheet data (RACM + Issues) for multi-tab view, filtered by active audit."""
    audit_id = get_active_audit_id()
    user = get_current_user()

    if audit_id:
        # Get audit-filtered data
        risks = db.get_risks_by_audit(audit_id)
        issues = db.get_issues_by_audit(audit_id)
        audit = db.get_audit(audit_id)

        # Build audit context for permission checks
        audit_context = {
            'id': audit_id,
            'auditor_id': audit.get('auditor_id') if audit else None,
            'reviewer_id': audit.get('reviewer_id') if audit else None
        }

        # Convert to spreadsheet format
        # Columns: Risk ID, Risk, Control ID, Control Owner, DE Testing, DE Conclusion,
        #          OE Testing, OE Conclusion, Status, Flowchart, Task, Evidence
        racm_rows = []
        racm_metadata = []  # Additional data for workflow
        conn = db._get_conn()
        for r in risks:
            fc = conn.execute("SELECT name FROM flowcharts WHERE risk_id = ?", (r['id'],)).fetchone()
            task = conn.execute("SELECT title FROM tasks WHERE risk_id = ?", (r['id'],)).fetchone()
            racm_rows.append([
                r['risk_id'],                                          # 0: Risk ID
                r['risk'] or '',                                       # 1: Risk
                r['control_id'] or '',                                 # 2: Control ID
                r['control_owner'] or '',                              # 3: Control Owner
                r.get('design_effectiveness_testing') or '',           # 4: DE Testing
                r.get('design_effectiveness_conclusion') or '',        # 5: DE Conclusion
                r.get('operational_effectiveness_test') or '',         # 6: OE Testing
                r.get('operational_effectiveness_conclusion') or '',   # 7: OE Conclusion
                r['status'] or '',                                     # 8: Status
                fc['name'] if fc else '',                              # 9: Flowchart
                task['title'] if task else '',                         # 10: Task
                ''                                                     # 11: Evidence (read-only)
            ])
            # Add workflow metadata for each row
            record_status = r.get('record_status') or 'draft'
            permissions = get_record_permissions(user, audit_context, r)
            racm_metadata.append({
                'id': r['id'],
                'record_status': record_status,
                'current_owner_role': r.get('current_owner_role') or 'auditor',
                'signed_off_by': r.get('signed_off_by'),
                'signed_off_at': r.get('signed_off_at'),
                'admin_lock_reason': r.get('admin_lock_reason'),
                'permissions': permissions
            })
        conn.close()

        issues_rows = []
        issues_metadata = []
        for issue in issues:
            issues_rows.append([
                issue['issue_id'],
                issue['risk_id'] or '',
                issue['title'],
                issue['description'] or '',
                issue['severity'],
                issue['status'],
                issue['assigned_to'] or '',
                issue['due_date'] or ''
            ])
            record_status = issue.get('record_status') or 'draft'
            permissions = get_record_permissions(user, audit_context, issue)
            issues_metadata.append({
                'id': issue['id'],
                'record_status': record_status,
                'current_owner_role': issue.get('current_owner_role') or 'auditor',
                'signed_off_by': issue.get('signed_off_by'),
                'signed_off_at': issue.get('signed_off_at'),
                'admin_lock_reason': issue.get('admin_lock_reason'),
                'permissions': permissions
            })

        # Get audit team for reviewer selection
        audit_team = db.get_audit_team(audit_id) if audit_id else []
        audit_auditors = [m for m in audit_team if m['team_role'] == 'auditor']
        audit_reviewers = [m for m in audit_team if m['team_role'] == 'reviewer']

        return jsonify({
            'racm': racm_rows,
            'racm_metadata': racm_metadata,
            'issues': issues_rows,
            'issues_metadata': issues_metadata,
            'audit': {
                'id': audit_id,
                'auditor_id': audit.get('auditor_id') if audit else None,
                'reviewer_id': audit.get('reviewer_id') if audit else None,
                'title': audit.get('title') if audit else None
            },
            'audit_team': {
                'auditors': audit_auditors,
                'reviewers': audit_reviewers
            },
            'user_role': get_user_role(user)
        })
    else:
        # No audit selected - return empty data
        return jsonify({
            'racm': [],
            'issues': []
        })

@app.route('/api/data', methods=['POST'])
@require_login
def save_data():
    """Save all spreadsheet data from multi-tab view."""
    data = request.json
    audit_id = get_active_audit_id()

    # Handle both old format (array) and new format (object with racm/issues)
    if isinstance(data, list):
        # Old format - just RACM data
        db.save_from_spreadsheet(data)
    else:
        # New format - both sheets
        if 'racm' in data:
            db.save_from_spreadsheet(data['racm'])
        if 'issues' in data:
            db.save_issues_from_spreadsheet(data['issues'])

    # Associate all risks and issues without audit_id to the active audit
    if audit_id:
        conn = db._get_conn()
        conn.execute("UPDATE risks SET audit_id = ? WHERE audit_id IS NULL", (audit_id,))
        conn.execute("UPDATE issues SET audit_id = ? WHERE audit_id IS NULL", (audit_id,))
        conn.commit()
        conn.close()

    increment_data_version()
    return jsonify({'status': 'saved'})

# ==================== Risks API (for direct access) ====================

@app.route('/api/risks', methods=['GET'])
@require_login
def get_risks():
    """Get all risks as JSON."""
    return jsonify(db.get_all_risks())

@app.route('/api/risks/<risk_id>', methods=['GET'])
@require_login
def get_risk(risk_id):
    """Get a single risk."""
    risk = db.get_risk(risk_id)
    return jsonify(risk) if risk else not_found_response('Risk')

@app.route('/api/risks', methods=['POST'])
@require_login
@require_non_viewer
def create_risk():
    """Create a new risk."""
    data = request.json or {}

    # Validate required field
    risk_id = data.get('risk_id', '').strip() if data.get('risk_id') else ''
    if not risk_id:
        return jsonify({'error': 'risk_id is required'}), 400

    # Thread-safe data version update
    new_id = db.create_risk(
        risk_id=risk_id,
        risk=data.get('risk', data.get('risk_description', '')),
        control_id=data.get('control_id', data.get('control_description', '')),
        control_owner=data.get('control_owner', ''),
        status=data.get('status', 'Not Complete')
    )
    increment_data_version()
    return jsonify({'status': 'created', 'id': new_id})

@app.route('/api/risks/<risk_id>', methods=['PUT'])
@require_login
@require_non_viewer
def update_risk(risk_id):
    """Update a risk."""
    # Thread-safe data version update
    db.update_risk(risk_id, **request.json)
    increment_data_version()
    return jsonify({'status': 'updated'})

@app.route('/api/risks/<risk_id>', methods=['DELETE'])
@require_login
@require_non_viewer
def delete_risk(risk_id):
    """Delete a risk."""
    if db.delete_risk(risk_id):
        increment_data_version()
        return jsonify({'status': 'deleted'})
    return not_found_response('Risk')

# ==================== Workflow State Transition API ====================

@app.route('/api/records/<record_type>/<int:record_id>/submit-for-review', methods=['POST'])
@require_login
def submit_for_review(record_type, record_id):
    """Submit a record for review (draft -> in_review).

    Requires reviewer_id in request body - auditor must select which reviewer.
    """
    if record_type not in ('risk', 'issue'):
        return jsonify({'error': 'Invalid record type'}), 400

    user = get_current_user()
    data = request.get_json() or {}

    # Get the record with audit info
    record = db.get_record_with_audit(record_type, record_id)
    if not record:
        return jsonify({'error': 'Record not found'}), 404

    audit = {
        'id': record.get('audit_id'),
        'auditor_id': record.get('auditor_id'),
        'reviewer_id': record.get('reviewer_id')
    }

    # Check if user can submit for review
    if not can_transition_record(user, audit, record, 'submit_for_review'):
        return jsonify({'error': 'Cannot submit this record for review'}), 403

    # Require reviewer_id selection
    reviewer_id = data.get('reviewer_id')
    if not reviewer_id:
        return jsonify({'error': 'reviewer_id is required - you must select a reviewer'}), 400

    # Validate that the selected reviewer is actually assigned to this audit
    if not db.is_reviewer_on_audit(reviewer_id, audit['id']):
        return jsonify({'error': 'Selected reviewer is not assigned to this audit'}), 400

    notes = data.get('notes', '')

    # Perform the transition with assigned_reviewer_id
    db.update_record_status(
        record_type, record_id,
        new_status='in_review',
        new_owner_role='reviewer',
        user_id=user['id'],
        assigned_reviewer_id=reviewer_id
    )

    # Log the state change
    db.create_state_history(
        record_type, record_id,
        from_status=record.get('record_status', 'draft'),
        to_status='in_review',
        action='submit_for_review',
        performed_by=user['id'],
        notes=notes
    )

    increment_data_version()
    return jsonify({'status': 'ok', 'record_status': 'in_review', 'assigned_reviewer_id': reviewer_id})


@app.route('/api/records/<record_type>/<int:record_id>/return-to-auditor', methods=['POST'])
@require_login
def return_to_auditor(record_type, record_id):
    """Return a record to the auditor (in_review -> draft)."""
    if record_type not in ('risk', 'issue'):
        return jsonify({'error': 'Invalid record type'}), 400

    user = get_current_user()
    record = db.get_record_with_audit(record_type, record_id)
    if not record:
        return jsonify({'error': 'Record not found'}), 404

    audit = {
        'id': record.get('audit_id'),
        'auditor_id': record.get('auditor_id'),
        'reviewer_id': record.get('reviewer_id')
    }

    if not can_transition_record(user, audit, record, 'return_to_auditor'):
        return jsonify({'error': 'Cannot return this record'}), 403

    data = request.get_json() or {}
    notes = data.get('notes', '')

    # Notes are required for return
    if not notes.strip():
        return jsonify({'error': 'Feedback notes are required when returning a record'}), 400

    # Perform the transition
    db.update_record_status(
        record_type, record_id,
        new_status='draft',
        new_owner_role='auditor',
        user_id=user['id']
    )

    # Log the state change
    db.create_state_history(
        record_type, record_id,
        from_status='in_review',
        to_status='draft',
        action='return_to_auditor',
        performed_by=user['id'],
        notes=notes
    )

    increment_data_version()
    return jsonify({'status': 'ok', 'record_status': 'draft'})


@app.route('/api/records/<record_type>/<int:record_id>/sign-off', methods=['POST'])
@require_login
def sign_off_record(record_type, record_id):
    """Sign off a record (in_review -> signed_off)."""
    if record_type not in ('risk', 'issue'):
        return jsonify({'error': 'Invalid record type'}), 400

    user = get_current_user()
    record = db.get_record_with_audit(record_type, record_id)
    if not record:
        return jsonify({'error': 'Record not found'}), 404

    audit = {
        'id': record.get('audit_id'),
        'auditor_id': record.get('auditor_id'),
        'reviewer_id': record.get('reviewer_id')
    }

    if not can_transition_record(user, audit, record, 'sign_off'):
        return jsonify({'error': 'Cannot sign off this record'}), 403

    data = request.get_json() or {}
    confirmation = data.get('confirmation', '')

    # Require typed confirmation
    if confirmation != 'SIGN OFF':
        return jsonify({'error': 'Invalid confirmation. Type "SIGN OFF" to confirm.'}), 400

    # Perform the transition
    db.update_record_status(
        record_type, record_id,
        new_status='signed_off',
        new_owner_role='none',
        user_id=user['id'],
        signed_off_by=user['id'],
        signed_off_at=datetime.now().isoformat()
    )

    # Log the state change
    db.create_state_history(
        record_type, record_id,
        from_status='in_review',
        to_status='signed_off',
        action='sign_off',
        performed_by=user['id']
    )

    increment_data_version()
    return jsonify({'status': 'ok', 'record_status': 'signed_off'})


@app.route('/api/admin/records/<record_type>/<int:record_id>/lock', methods=['POST'])
@require_login
@require_admin
def admin_lock_record(record_type, record_id):
    """Admin: Lock a record (any -> admin_hold)."""
    if record_type not in ('risk', 'issue'):
        return jsonify({'error': 'Invalid record type'}), 400

    user = get_current_user()
    record = db.get_record_with_audit(record_type, record_id)
    if not record:
        return jsonify({'error': 'Record not found'}), 404

    current_status = record.get('record_status', 'draft')
    if current_status == 'admin_hold':
        return jsonify({'error': 'Record is already on admin hold'}), 400

    data = request.get_json() or {}
    reason = data.get('reason', '')

    # Reason is required for admin lock
    if not reason.strip():
        return jsonify({'error': 'Reason is required for admin lock'}), 400

    # Perform the transition
    db.update_record_status(
        record_type, record_id,
        new_status='admin_hold',
        new_owner_role='none',
        user_id=user['id'],
        admin_lock_reason=reason,
        admin_locked_by=user['id'],
        admin_locked_at=datetime.now().isoformat()
    )

    # Log the state change
    db.create_state_history(
        record_type, record_id,
        from_status=current_status,
        to_status='admin_hold',
        action='admin_lock',
        performed_by=user['id'],
        reason=reason
    )

    increment_data_version()
    return jsonify({'status': 'ok', 'record_status': 'admin_hold'})


@app.route('/api/admin/records/<record_type>/<int:record_id>/unlock', methods=['POST'])
@require_login
@require_admin
def admin_unlock_record(record_type, record_id):
    """Admin: Unlock a record from admin_hold."""
    if record_type not in ('risk', 'issue'):
        return jsonify({'error': 'Invalid record type'}), 400

    user = get_current_user()
    record = db.get_record_with_audit(record_type, record_id)
    if not record:
        return jsonify({'error': 'Record not found'}), 404

    current_status = record.get('record_status', 'draft')
    if current_status != 'admin_hold':
        return jsonify({'error': 'Record is not on admin hold'}), 400

    data = request.get_json() or {}
    reason = data.get('reason', '')
    return_to = data.get('return_to', 'draft')

    if not reason.strip():
        return jsonify({'error': 'Reason is required for unlock'}), 400

    if return_to not in ('draft', 'in_review'):
        return jsonify({'error': 'return_to must be "draft" or "in_review"'}), 400

    new_owner_role = 'auditor' if return_to == 'draft' else 'reviewer'

    # Perform the transition
    db.update_record_status(
        record_type, record_id,
        new_status=return_to,
        new_owner_role=new_owner_role,
        user_id=user['id'],
        admin_lock_reason=None,
        admin_locked_by=None,
        admin_locked_at=None
    )

    # Log the state change
    db.create_state_history(
        record_type, record_id,
        from_status='admin_hold',
        to_status=return_to,
        action='admin_unlock',
        performed_by=user['id'],
        reason=reason
    )

    increment_data_version()
    return jsonify({'status': 'ok', 'record_status': return_to})


@app.route('/api/admin/records/<record_type>/<int:record_id>/unlock-signoff', methods=['POST'])
@require_login
@require_admin
def admin_unlock_signoff(record_type, record_id):
    """Admin: Unlock a signed-off record (requires stronger confirmation)."""
    if record_type not in ('risk', 'issue'):
        return jsonify({'error': 'Invalid record type'}), 400

    user = get_current_user()
    record = db.get_record_with_audit(record_type, record_id)
    if not record:
        return jsonify({'error': 'Record not found'}), 404

    current_status = record.get('record_status', 'draft')
    if current_status != 'signed_off':
        return jsonify({'error': 'Record is not signed off'}), 400

    data = request.get_json() or {}
    reason = data.get('reason', '')
    return_to = data.get('return_to', 'draft')
    confirmation = data.get('confirmation', '')

    if not reason.strip():
        return jsonify({'error': 'Reason is required'}), 400

    if return_to not in ('draft', 'in_review'):
        return jsonify({'error': 'return_to must be "draft" or "in_review"'}), 400

    if confirmation != 'UNLOCK SIGNED OFF':
        return jsonify({'error': 'Invalid confirmation. Type "UNLOCK SIGNED OFF" to confirm.'}), 400

    new_owner_role = 'auditor' if return_to == 'draft' else 'reviewer'

    # Perform the transition
    db.update_record_status(
        record_type, record_id,
        new_status=return_to,
        new_owner_role=new_owner_role,
        user_id=user['id'],
        signed_off_by=None,
        signed_off_at=None
    )

    # Log the state change
    db.create_state_history(
        record_type, record_id,
        from_status='signed_off',
        to_status=return_to,
        action='admin_unlock_signoff',
        performed_by=user['id'],
        reason=reason
    )

    increment_data_version()
    return jsonify({'status': 'ok', 'record_status': return_to})


@app.route('/api/records/<record_type>/<int:record_id>/history', methods=['GET'])
@require_login
def get_record_history(record_type, record_id):
    """Get the state transition history for a record."""
    if record_type not in ('risk', 'issue'):
        return jsonify({'error': 'Invalid record type'}), 400

    history = db.get_record_history(record_type, record_id)
    return jsonify(history)


@app.route('/api/records/<record_type>/<int:record_id>/permissions', methods=['GET'])
@require_login
def get_permissions_for_record(record_type, record_id):
    """Get the current user's permissions for a specific record."""
    if record_type not in ('risk', 'issue'):
        return jsonify({'error': 'Invalid record type'}), 400

    user = get_current_user()
    record = db.get_record_with_audit(record_type, record_id)
    if not record:
        return jsonify({'error': 'Record not found'}), 404

    audit = {
        'id': record.get('audit_id'),
        'auditor_id': record.get('auditor_id'),
        'reviewer_id': record.get('reviewer_id')
    }

    permissions = get_record_permissions(user, audit, record)
    return jsonify(permissions)


@app.route('/api/audits/<int:audit_id>/workflow-summary', methods=['GET'])
@require_login
def get_audit_workflow_summary(audit_id):
    """Get workflow status summary for an audit."""
    summary = db.get_workflow_summary(audit_id)
    return jsonify(summary)


# ==================== Viewer Management API ====================

@app.route('/api/audits/<int:audit_id>/viewers', methods=['GET'])
@require_login
def get_audit_viewers(audit_id):
    """Get all viewers for an audit."""
    viewers = db.get_audit_viewers(audit_id)
    return jsonify(viewers)


@app.route('/api/audits/<int:audit_id>/viewers', methods=['POST'])
@require_login
def add_audit_viewer(audit_id):
    """Add a viewer to an audit."""
    user = get_current_user()
    audit = db.get_audit(audit_id)
    if not audit:
        return jsonify({'error': 'Audit not found'}), 404

    # Only auditor, reviewer of this audit, or admin can add viewers
    is_assigned = (audit.get('auditor_id') == user['id'] or
                   audit.get('reviewer_id') == user['id'])
    if not user.get('is_admin') and not is_assigned:
        return jsonify({'error': 'Not authorized to manage viewers'}), 403

    data = request.get_json() or {}
    viewer_user_id = data.get('user_id')
    if not viewer_user_id:
        return jsonify({'error': 'user_id is required'}), 400

    viewer_id = db.add_audit_viewer(audit_id, viewer_user_id, user['id'])
    return jsonify({'status': 'ok', 'viewer_id': viewer_id})


@app.route('/api/audits/<int:audit_id>/viewers/<int:viewer_user_id>', methods=['DELETE'])
@require_login
def remove_audit_viewer(audit_id, viewer_user_id):
    """Remove a viewer from an audit."""
    user = get_current_user()
    audit = db.get_audit(audit_id)
    if not audit:
        return jsonify({'error': 'Audit not found'}), 404

    # Only auditor, reviewer of this audit, or admin can remove viewers
    is_assigned = (audit.get('auditor_id') == user['id'] or
                   audit.get('reviewer_id') == user['id'])
    if not user.get('is_admin') and not is_assigned:
        return jsonify({'error': 'Not authorized to manage viewers'}), 403

    if db.remove_audit_viewer(audit_id, viewer_user_id):
        return jsonify({'status': 'ok'})
    return jsonify({'error': 'Viewer not found'}), 404


# ==================== Audit Assignment API ====================

@app.route('/api/audits/<int:audit_id>/assignment', methods=['PUT'])
@require_login
@require_admin
def update_audit_assignment(audit_id):
    """Update auditor/reviewer assignment for an audit."""
    data = request.get_json() or {}
    auditor_id = data.get('auditor_id')
    reviewer_id = data.get('reviewer_id')

    if db.update_audit_assignment(audit_id, auditor_id, reviewer_id):
        return jsonify({'status': 'ok'})
    return jsonify({'error': 'No changes made'}), 400


# ==================== Admin Audit Log ====================

@app.route('/api/admin/audit-log', methods=['GET'])
@require_login
@require_admin
def get_admin_audit_log():
    """Get all state history entries (admin audit log)."""
    filters = {
        'from_date': request.args.get('from_date'),
        'to_date': request.args.get('to_date'),
        'user_id': request.args.get('user_id', type=int),
        'action': request.args.get('action'),
        'record_type': request.args.get('record_type')
    }
    # Remove None values
    filters = {k: v for k, v in filters.items() if v is not None}

    history = db.get_all_state_history(filters if filters else None)
    return jsonify(history)


@app.route('/api/admin/records-in-hold', methods=['GET'])
@require_login
@require_admin
def get_records_in_admin_hold():
    """Get all records in admin_hold or signed_off status for admin override panel."""
    # Get admin_hold records (already returns combined list with record_type)
    hold_data = db.get_records_in_admin_hold()
    hold_records = hold_data.get('risks', []) + hold_data.get('issues', [])
    # Mark record types
    for r in hold_records:
        if 'risk_id' in r and r.get('risk_id'):
            r['record_type'] = 'risk'
        else:
            r['record_type'] = 'issue'

    # Get signed_off records
    signedoff_records = db.get_all_records_by_status('signed_off')

    return jsonify(hold_records + signedoff_records)


# ==================== Flowchart API ====================

@app.route('/flowchart')
@app.route('/flowchart/<flowchart_id>')
@require_login
def flowchart(flowchart_id=None):
    ctx = get_audit_context()
    is_admin = session.get('is_admin', False)
    return render_template('flowchart.html', flowchart_id=flowchart_id, active_page='flowchart',
                           is_admin=is_admin, **ctx)

@app.route('/api/flowchart/<flowchart_id>', methods=['GET'])
@require_login
def get_flowchart(flowchart_id):
    """Get a flowchart by name, including permissions for the linked risk."""
    fc = db.get_flowchart(flowchart_id)
    if not fc:
        return jsonify(None)

    # Get permissions for the linked risk
    permissions = {'canEdit': False}
    if fc.get('risk_id'):
        # Get the risk record and audit context
        conn = db._get_conn()
        risk_row = conn.execute("SELECT * FROM risks WHERE id = ?", (fc['risk_id'],)).fetchone()
        conn.close()

        if risk_row:
            risk = dict(risk_row)
            user = get_current_user()
            audit_id = risk.get('audit_id') or get_active_audit_id()
            audit = db.get_audit(audit_id) if audit_id else None
            audit_context = {
                'id': audit_id,
                'auditor_id': audit.get('auditor_id') if audit else None,
                'reviewer_id': audit.get('reviewer_id') if audit else None
            }
            permissions = get_record_permissions(user, audit_context, risk)

    return jsonify({
        'data': fc['data'],
        'permissions': permissions
    })

@app.route('/api/flowchart/<flowchart_id>', methods=['POST'])
@require_login
def save_flowchart(flowchart_id):
    """Save a flowchart. Enforces edit permissions for the linked risk."""
    # Check permissions for the linked risk
    fc = db.get_flowchart(flowchart_id)
    if fc and fc.get('risk_id'):
        conn = db._get_conn()
        risk_row = conn.execute("SELECT * FROM risks WHERE id = ?", (fc['risk_id'],)).fetchone()
        conn.close()

        if risk_row:
            risk = dict(risk_row)
            user = get_current_user()
            audit_id = risk.get('audit_id') or get_active_audit_id()
            audit = db.get_audit(audit_id) if audit_id else None
            audit_context = {
                'id': audit_id,
                'auditor_id': audit.get('auditor_id') if audit else None,
                'reviewer_id': audit.get('reviewer_id') if audit else None
            }
            if not can_edit_record(user, audit_context, risk):
                return jsonify({'error': 'Cannot edit flowchart - record is not editable'}), 403

    db.save_flowchart(flowchart_id, request.json)
    increment_data_version()
    return jsonify({'status': 'saved'})

@app.route('/api/flowcharts', methods=['GET'])
@require_login
def list_flowcharts():
    """List flowchart names for the active audit."""
    audit_id = get_active_audit_id()
    if audit_id:
        flowcharts = db.get_flowcharts_by_audit(audit_id)
    else:
        flowcharts = db.get_all_flowcharts()
    return jsonify([f['name'] for f in flowcharts])

# ==================== Test Documents API ====================

@app.route('/api/test-document/<risk_code>/<doc_type>', methods=['GET'])
def get_test_document(risk_code, doc_type):
    """Get a test document by risk code and type (de_testing or oe_testing)."""
    if doc_type not in ('de_testing', 'oe_testing'):
        return jsonify({'error': 'Invalid document type'}), 400
    doc = db.get_test_document_by_risk_code(risk_code, doc_type)
    if doc:
        return jsonify({'content': doc['content'], 'id': doc['id']})
    return jsonify({'content': '', 'id': None})

@app.route('/api/test-document/<risk_code>/<doc_type>', methods=['POST'])
@require_login
def save_test_document(risk_code, doc_type):
    """Save a test document. Enforces edit permissions for the risk."""
    if doc_type not in ('de_testing', 'oe_testing'):
        return jsonify({'error': 'Invalid document type'}), 400

    # Get the risk and check permissions
    risk = db.get_risk(risk_code)
    if not risk:
        return jsonify({'error': 'Risk not found'}), 404

    user = get_current_user()
    # Admins can always edit
    if not user.get('is_admin'):
        audit_id = risk.get('audit_id') or get_active_audit_id()
        audit = db.get_audit(audit_id) if audit_id else None
        audit_context = {
            'id': audit_id,
            'auditor_id': audit.get('auditor_id') if audit else None,
            'reviewer_id': audit.get('reviewer_id') if audit else None
        }
        if not can_edit_record(user, audit_context, risk):
            return jsonify({'error': 'Cannot edit test document - record is not editable'}), 403

    data = request.json
    content = data.get('content', '')
    doc_id = db.save_test_document_by_risk_code(risk_code, doc_type, content)
    increment_data_version()
    return jsonify({'status': 'saved', 'id': doc_id})

@app.route('/api/test-document/<risk_code>/<doc_type>/exists', methods=['GET'])
def test_document_exists(risk_code, doc_type):
    """Check if a test document exists."""
    if doc_type not in ('de_testing', 'oe_testing'):
        return jsonify({'error': 'Invalid document type'}), 400
    exists = db.has_test_document(risk_code, doc_type)
    return jsonify({'exists': exists})

# ==================== Kanban API ====================

@app.route('/kanban')
@app.route('/kanban/<board_id>')
@require_login
def kanban(board_id='default'):
    ctx = get_audit_context()
    return render_template('kanban.html', board_id=board_id, active_page='kanban', **ctx)


@app.route('/audit-plan')
@require_login
def audit_plan():
    """Annual Audit Plan page with spreadsheet and kanban views."""
    ctx = get_audit_context()
    user = get_current_user()
    user_role = get_user_role(user) if user else None
    is_admin = session.get('is_admin', False)
    return render_template('audit_plan.html', active_page='audit-plan',
                           user_role=user_role, is_admin=is_admin, **ctx)


# ==================== Annual Audit Plan API ====================

@app.route('/api/audits', methods=['GET'])
@require_login
def get_audits():
    """Get all audits from the annual audit plan.

    Everyone can see the full audit plan (it's the annual overview).
    Access restrictions apply when selecting/working on specific audits.
    """
    audits = db.get_all_audits()
    return jsonify(audits)


@app.route('/api/audits', methods=['POST'])
@require_login
@require_non_viewer
def create_audit():
    """Create a new audit in the annual plan."""
    data = request.json
    title = data.get('title', '').strip()
    if not title:
        return jsonify({'error': 'Title is required'}), 400

    audit_id = db.create_audit(
        title=title,
        description=data.get('description'),
        audit_area=data.get('audit_area'),
        owner=data.get('owner'),
        planned_start=data.get('planned_start'),
        planned_end=data.get('planned_end'),
        quarter=data.get('quarter'),
        status=data.get('status', 'planning'),
        priority=data.get('priority', 'medium'),
        risk_rating=data.get('risk_rating'),
        estimated_hours=data.get('estimated_hours'),
        notes=data.get('notes')
    )
    increment_data_version()
    return jsonify({'id': audit_id, 'status': 'created'})


@app.route('/api/audits/<int:audit_id>', methods=['GET'])
@require_login
def get_audit(audit_id):
    """Get a single audit by ID."""
    audit = db.get_audit(audit_id)
    if not audit:
        return not_found_response('Audit')
    return jsonify(audit)


@app.route('/api/audits/<int:audit_id>', methods=['PUT'])
@require_login
@require_non_viewer
def update_audit(audit_id):
    """Update an audit. Only updates fields present in the request."""
    audit = db.get_audit(audit_id)
    if not audit:
        return not_found_response('Audit')

    data = request.json
    # Only pass fields that are actually in the request (partial update)
    allowed_fields = {'title', 'description', 'audit_area', 'owner', 'planned_start',
                      'planned_end', 'actual_start', 'actual_end', 'quarter', 'status',
                      'priority', 'estimated_hours', 'actual_hours', 'risk_rating', 'notes'}
    update_data = {k: v for k, v in data.items() if k in allowed_fields}

    if update_data:
        db.update_audit(audit_id, **update_data)
        increment_data_version()
    return jsonify({'status': 'updated'})


@app.route('/api/audits/<int:audit_id>', methods=['DELETE'])
@require_login
@require_non_viewer
def delete_audit(audit_id):
    """Delete an audit from the annual plan."""
    deleted = db.delete_audit(audit_id)
    if not deleted:
        return not_found_response('Audit')
    increment_data_version()
    return jsonify({'status': 'deleted'})


@app.route('/api/audits/spreadsheet', methods=['GET'])
@require_login
def get_audits_spreadsheet():
    """Get audits in spreadsheet format (array of arrays)."""
    data = db.get_audits_as_spreadsheet()
    return jsonify(data)


@app.route('/api/audits/spreadsheet', methods=['POST'])
@require_login
@require_non_viewer
def save_audits_spreadsheet():
    """Save audits from spreadsheet format."""
    data = request.json
    if not isinstance(data, list):
        return jsonify({'error': 'Expected array of arrays'}), 400
    result = db.save_audits_from_spreadsheet(data)
    increment_data_version()
    return jsonify({'status': 'saved', **result})


@app.route('/api/audits/kanban', methods=['GET'])
@require_login
def get_audits_kanban():
    """Get audits in kanban board format."""
    board = db.get_audits_as_kanban()
    return jsonify(board)


@app.route('/api/audits/summary', methods=['GET'])
@require_login
def get_audits_summary():
    """Get summary statistics for the annual audit plan."""
    summary = db.get_audit_summary()
    return jsonify(summary)


@app.route('/library')
@require_login
def library():
    """Audit Library page for managing reference documents."""
    ctx = get_audit_context()
    return render_template('library.html', active_page='library', **ctx)


# ==================== LIBRARY API ====================

@app.route('/api/library/documents', methods=['GET'])
@require_login
def list_library_documents():
    """List all library documents."""
    doc_type = request.args.get('type')
    documents = db.list_library_documents(doc_type)
    return jsonify(documents)


@app.route('/api/library/documents', methods=['POST'])
@require_login
@require_non_viewer
def upload_library_document():
    """Upload a new library document."""
    if 'file' not in request.files:
        return error_response('No file provided')

    file = request.files['file']
    if file.filename == '':
        return error_response('No file selected')

    if not allowed_file(file.filename):
        return error_response('File type not allowed')

    # Get metadata from form
    name = request.form.get('name', file.filename.rsplit('.', 1)[0])
    doc_type = request.form.get('doc_type', 'framework')
    source = request.form.get('source', '')
    description = request.form.get('description', '')

    # Save file
    filename = secure_filename(file.filename)
    timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
    unique_filename = f"{timestamp}_{filename}"
    filepath = os.path.join(LIBRARY_FOLDER, unique_filename)
    file.save(filepath)

    # Get file info
    file_size = os.path.getsize(filepath)
    mime_type = file.content_type

    # Create database record
    doc_id = db.add_library_document(
        name=name,
        filename=unique_filename,
        original_filename=filename,
        doc_type=doc_type,
        source=source,
        description=description,
        file_size=file_size,
        mime_type=mime_type
    )

    # Process document in background (chunking and embedding)
    # For now, do it synchronously - can be made async later
    try:
        chunk_count = process_library_document(filepath, doc_id, mime_type)
        return jsonify({
            'id': doc_id,
            'name': name,
            'chunks': chunk_count,
            'message': f'Document uploaded and processed into {chunk_count} chunks'
        })
    except Exception as e:
        logger.error(f"[Library] Error processing document: {e}")
        return jsonify({
            'id': doc_id,
            'name': name,
            'chunks': 0,
            'warning': f'Document uploaded but processing failed: {str(e)}'
        })


@app.route('/api/library/documents/<int:doc_id>', methods=['GET'])
@require_login
def get_library_document(doc_id):
    """Get a library document's details."""
    doc = db.get_library_document(doc_id)
    if not doc:
        return not_found_response('Document')
    return jsonify(doc)


@app.route('/api/library/documents/<int:doc_id>', methods=['PUT'])
@require_login
@require_non_viewer
def update_library_document_route(doc_id):
    """Update a library document's metadata."""
    data = request.json
    if not data:
        return error_response('No data provided')

    success = db.update_library_document(doc_id, **data)
    if success:
        return jsonify({'status': 'updated'})
    return error_response('Update failed')


@app.route('/api/library/documents/<int:doc_id>', methods=['DELETE'])
@require_login
@require_non_viewer
def delete_library_document_route(doc_id):
    """Delete a library document and its chunks."""
    doc = db.get_library_document(doc_id)
    if not doc:
        return not_found_response('Document')

    # Delete file from disk
    filepath = os.path.join(LIBRARY_FOLDER, doc['filename'])
    if os.path.exists(filepath):
        os.remove(filepath)

    # Delete from database
    db.delete_library_document(doc_id)
    return jsonify({'status': 'deleted'})


@app.route('/api/library/documents/<int:doc_id>/file', methods=['GET'])
@require_login
def download_library_document(doc_id):
    """Download/view the original library document file."""
    doc = db.get_library_document(doc_id)
    if not doc:
        return not_found_response('Document')

    filepath = os.path.join(LIBRARY_FOLDER, doc['filename'])
    if not os.path.exists(filepath):
        return error_response('File not found on disk')

    # Determine if we should display inline (view) or download
    disposition = request.args.get('disposition', 'inline')

    return send_file(
        filepath,
        mimetype=doc.get('mime_type', 'application/octet-stream'),
        as_attachment=(disposition == 'attachment'),
        download_name=doc.get('original_filename', doc['filename'])
    )


@app.route('/api/library/documents/<int:doc_id>/chunks', methods=['GET'])
@require_login
def get_library_document_chunks(doc_id):
    """Get all chunks for a document."""
    chunks = db.get_library_chunks(doc_id)
    return jsonify(chunks)


@app.route('/api/library/search', methods=['POST'])
@require_login
def search_library():
    """Search the library using semantic search."""
    data = request.json
    query = data.get('query', '')
    limit = data.get('limit', 5)

    if not query:
        return error_response('No query provided')

    # Generate embedding for query
    query_embedding = generate_embedding(query)

    if query_embedding:
        # Vector search
        results = db.search_library(query_embedding, limit)
    else:
        # Fallback to keyword search
        results = db.search_library_keyword(query, limit)

    return jsonify({
        'query': query,
        'results': results,
        'method': 'vector' if query_embedding else 'keyword'
    })


@app.route('/api/library/stats', methods=['GET'])
@require_login
def get_library_stats():
    """Get library statistics."""
    stats = db.get_library_stats()
    return jsonify(stats)


@app.route('/api/kanban/<board_id>', methods=['GET'])
@require_login
def get_kanban(board_id):
    """Get kanban board in legacy format."""
    kanban_format = db.get_kanban_format()
    return jsonify(kanban_format['boards'].get(board_id))

@app.route('/api/kanban/<board_id>', methods=['POST'])
@require_login
@require_non_viewer
def save_kanban(board_id):
    """Save kanban board from legacy format."""
    # Thread-safe data version update
    data = request.json

    # Clear existing tasks and recreate from the board data
    # This is a simple approach - in production you'd want smarter sync
    for col in data.get('columns', []):
        for item in col.get('items', []):
            task = db.get_task(int(item['id'])) if item['id'].isdigit() else None
            if task:
                db.update_task(int(item['id']),
                              title=item.get('title', ''),
                              description=item.get('description', ''),
                              priority=item.get('priority', 'medium'),
                              assignee=item.get('assignee', ''),
                              column_id=col['id'])

    increment_data_version()
    return jsonify({'status': 'saved'})

@app.route('/api/kanban/<board_id>/task', methods=['POST'])
@require_login
@require_non_viewer
def create_kanban_task(board_id):
    """Create a new task on the kanban board."""
    # Thread-safe data version update
    data = request.json

    task_id = db.create_task(
        title=data.get('title', 'New Task'),
        description=data.get('description', ''),
        priority=data.get('priority', 'medium'),
        assignee=data.get('assignee', ''),
        column_id=data.get('column', 'planning'),
        risk_id=data.get('riskId', None)
    )

    increment_data_version()
    return jsonify({'status': 'created', 'taskId': str(task_id)})

@app.route('/api/kanban/<board_id>/task/<task_id>', methods=['GET'])
@require_login
def get_kanban_task(board_id, task_id):
    """Get a specific task."""
    task = db.get_task(int(task_id))
    if task:
        return jsonify({
            'id': str(task['id']),
            'title': task['title'],
            'description': task['description'],
            'priority': task['priority'],
            'assignee': task['assignee'],
            'column': task['column_id'],
            'riskId': task['linked_risk_id'] or ''
        })
    return jsonify(None)

@app.route('/api/kanban/<board_id>/task/<task_id>', methods=['PUT'])
@require_login
@require_non_viewer
def update_kanban_task(board_id, task_id):
    """Update a specific task."""
    # Thread-safe data version update
    data = request.json or {}

    # Only include fields that were actually provided
    updates = {}
    if 'title' in data:
        updates['title'] = data['title']
    if 'description' in data:
        updates['description'] = data['description']
    if 'priority' in data:
        updates['priority'] = data['priority']
    if 'assignee' in data:
        updates['assignee'] = data['assignee']
    if 'column' in data:
        updates['column_id'] = data['column']
    if 'column_id' in data:
        updates['column_id'] = data['column_id']

    if updates:
        db.update_task(int(task_id), **updates)
        increment_data_version()
    return jsonify({'status': 'updated'})

# ==================== Tasks API (direct access) ====================

@app.route('/api/tasks', methods=['GET'])
@require_login
def get_tasks():
    """Get tasks filtered by active audit."""
    audit_id = get_active_audit_id()
    if audit_id:
        return jsonify(db.get_tasks_by_audit(audit_id))
    else:
        return jsonify([])

@app.route('/api/tasks/<int:task_id>', methods=['GET'])
def get_task(task_id):
    """Get a single task."""
    task = db.get_task(task_id)
    return jsonify(task) if task else not_found_response('Task')

@app.route('/api/tasks', methods=['POST'])
@require_login
def create_task():
    """Create a new task and associate with active audit."""
    data = request.json
    task_id = db.create_task(
        title=data.get('title'),
        description=data.get('description', ''),
        priority=data.get('priority', 'medium'),
        assignee=data.get('assignee', ''),
        column_id=data.get('column_id', 'planning'),
        risk_id=data.get('risk_id')
    )

    # Associate with active audit
    audit_id = get_active_audit_id()
    if audit_id:
        conn = db._get_conn()
        conn.execute("UPDATE tasks SET audit_id = ? WHERE id = ?", (audit_id, task_id))
        conn.commit()
        conn.close()

    increment_data_version()
    return jsonify({'status': 'created', 'id': task_id})

# ==================== AI Query API ====================

@app.route('/api/query', methods=['POST'])
@require_admin
def query_database():
    """Execute a read-only SQL query (for AI integration). Admin only."""
    data = request.json
    sql = data.get('sql', '')
    try:
        results = db.execute_query(sql)
        return jsonify({'results': results})
    except ValueError as e:
        # Validation errors (forbidden keywords) - safe to return
        return jsonify({'error': str(e)}), 400
    except Exception as e:
        # Log the full error but return generic message
        app.logger.error(f"SQL query error: {e}")
        return jsonify({'error': 'Query execution failed'}), 400

@app.route('/api/schema', methods=['GET'])
@require_login
def get_schema():
    """Get database schema for AI context."""
    return jsonify({'schema': db.get_schema()})

@app.route('/api/context', methods=['GET'])
@require_login
def get_context():
    """Get full database context for AI.

    For non-admins, returns only data from audits they have access to.
    """
    user = get_current_user()
    if user['is_admin']:
        return jsonify(db.get_full_context())

    # Non-admin: return scoped context
    audit_ids = db.get_user_audit_ids(user['id'])
    return jsonify(db.get_context_for_audits(audit_ids))

# ==================== Issues API ====================

@app.route('/api/issues', methods=['GET'])
@require_login
def get_issues():
    """Get all issues."""
    return jsonify(db.get_all_issues())

@app.route('/api/issues/<issue_id>', methods=['GET'])
@require_login
def get_issue(issue_id):
    """Get a single issue."""
    issue = db.get_issue(issue_id)
    return jsonify(issue) if issue else not_found_response('Issue')

@app.route('/api/issues', methods=['POST'])
@require_login
@require_non_viewer
def create_issue():
    """Create a new issue."""
    # Thread-safe data version update
    data = request.json
    issue_id = db.create_issue(
        risk_id=data.get('risk_id', ''),
        title=data.get('title', ''),
        description=data.get('description', ''),
        severity=data.get('severity', 'Medium'),
        status=data.get('status', 'Open'),
        assigned_to=data.get('assigned_to', ''),
        due_date=data.get('due_date')
    )
    increment_data_version()
    return jsonify({'status': 'created', 'issue_id': issue_id})

@app.route('/api/issues/<issue_id>', methods=['PUT'])
@require_login
@require_non_viewer
def update_issue(issue_id):
    """Update an issue."""
    if db.update_issue(issue_id, **request.json):
        increment_data_version()
        return jsonify({'status': 'updated'})
    return not_found_response('Issue')

@app.route('/api/issues/<issue_id>', methods=['DELETE'])
@require_login
@require_non_viewer
def delete_issue(issue_id):
    """Delete an issue."""
    if db.delete_issue(issue_id):
        increment_data_version()
        return jsonify({'status': 'deleted'})
    return not_found_response('Issue')

@app.route('/api/issues/<issue_id>/documentation', methods=['GET'])
@require_login
def get_issue_documentation(issue_id):
    """Get documentation for an issue."""
    doc = db.get_issue_documentation(issue_id)
    has_doc = db.has_issue_documentation(issue_id)
    return jsonify({'documentation': doc or '', 'has_documentation': has_doc})

@app.route('/api/issues/<issue_id>/documentation', methods=['POST'])
@require_login
def save_issue_documentation(issue_id):
    """Save documentation for an issue. Enforces edit permissions."""
    # Get the issue and check permissions
    issue = db.get_issue(issue_id)
    if not issue:
        return not_found_response('Issue')

    user = get_current_user()
    # Admins can always edit
    if not user.get('is_admin'):
        audit_id = issue.get('audit_id') or get_active_audit_id()
        audit = db.get_audit(audit_id) if audit_id else None
        audit_context = {
            'id': audit_id,
            'auditor_id': audit.get('auditor_id') if audit else None,
            'reviewer_id': audit.get('reviewer_id') if audit else None
        }
        if not can_edit_record(user, audit_context, issue):
            return jsonify({'error': 'Cannot edit issue documentation - record is not editable'}), 403

    data = request.json
    documentation = data.get('documentation', '')
    if db.save_issue_documentation(issue_id, documentation):
        increment_data_version()
        return jsonify({'status': 'saved'})
    return not_found_response('Issue')

@app.route('/api/issues/<issue_id>/documentation/exists', methods=['GET'])
@require_login
def issue_documentation_exists(issue_id):
    """Check if an issue has documentation."""
    return jsonify({'exists': db.has_issue_documentation(issue_id)})

# ==================== Issue Attachments API ====================

def _process_file_upload(file):
    """Process uploaded file and return file info dict.

    Returns dict with: filepath, original_filename, unique_filename, file_size, mime_type, extracted_text
    Raises ValueError on validation errors.
    Cleans up file on extraction failure.
    """
    if file.filename == '':
        raise ValueError('No file selected')

    if not allowed_file(file.filename):
        raise ValueError(f'File type not allowed. Allowed: {", ".join(ALLOWED_EXTENSIONS)}')

    # Generate unique filename
    original_filename = secure_filename(file.filename)
    ext = original_filename.rsplit('.', 1)[1].lower() if '.' in original_filename else ''
    unique_filename = f"{uuid.uuid4().hex}.{ext}" if ext else uuid.uuid4().hex

    # Save file
    filepath = os.path.join(app.config['UPLOAD_FOLDER'], unique_filename)
    file.save(filepath)

    try:
        file_size = os.path.getsize(filepath)
        mime_type = mimetypes.guess_type(original_filename)[0] or 'application/octet-stream'
        extracted_text = extract_text_from_file(filepath, mime_type)

        return {
            'filepath': filepath,
            'original_filename': original_filename,
            'unique_filename': unique_filename,
            'file_size': file_size,
            'mime_type': mime_type,
            'extracted_text': extracted_text
        }
    except Exception:
        # Clean up file on failure
        if os.path.exists(filepath):
            os.remove(filepath)
        raise


def _format_attachment_list(attachments: list, entity_type: str, entity_id: str) -> str:
    """Format attachment list for AI tool response."""
    if not attachments:
        return f"No file attachments found for {entity_type} {entity_id}."

    result = f"## File Attachments for {entity_type.title()} {entity_id}\n\n"
    for att in attachments:
        size_kb = att['file_size'] / 1024
        size_str = f"{size_kb:.1f} KB" if size_kb < 1024 else f"{size_kb/1024:.1f} MB"
        result += f"- **ID: {att['id']}** - {att['original_filename']} ({size_str})\n"
        result += f"  Type: {att['mime_type']}\n"
        if att.get('description'):
            result += f"  Description: {att['description']}\n"
        result += f"  Uploaded: {att['uploaded_at']}\n\n"

    return result


@app.route('/api/issues/<issue_id>/attachments', methods=['GET'])
@require_login
def get_issue_attachments(issue_id):
    """Get all attachments for an issue."""
    attachments = db.get_attachments_for_issue(issue_id)
    return jsonify(attachments)

@app.route('/api/issues/<issue_id>/attachments', methods=['POST'])
@require_login
def upload_issue_attachment(issue_id):
    """Upload a file attachment for an issue. Enforces edit permissions."""
    issue = db.get_issue(issue_id)
    if not issue:
        return not_found_response('Issue')

    allowed, error = check_attachment_permission(issue, 'issue')
    if not allowed:
        return error

    if 'file' not in request.files:
        return jsonify({'error': 'No file provided'}), 400

    try:
        file_info = _process_file_upload(request.files['file'])
    except ValueError as e:
        return jsonify({'error': str(e)}), 400

    attachment_id = db.add_attachment(
        issue_id=issue_id,
        filename=file_info['unique_filename'],
        original_filename=file_info['original_filename'],
        file_size=file_info['file_size'],
        mime_type=file_info['mime_type'],
        description=request.form.get('description', ''),
        extracted_text=file_info['extracted_text']
    )

    increment_data_version()
    return jsonify(attachment_upload_response(file_info, attachment_id))

@app.route('/api/attachments/<int:attachment_id>', methods=['GET'])
def download_attachment(attachment_id):
    """Download an attachment file."""
    attachment = db.get_attachment(attachment_id)
    if not attachment:
        return not_found_response('Attachment')

    return send_from_directory(
        app.config['UPLOAD_FOLDER'],
        attachment['filename'],
        download_name=attachment['original_filename'],
        as_attachment=True
    )

@app.route('/api/attachments/<int:attachment_id>', methods=['DELETE'])
@require_login
def delete_attachment(attachment_id):
    """Delete an issue attachment. Enforces edit permissions on the parent issue."""
    attachment = db.get_attachment(attachment_id)
    if not attachment:
        return not_found_response('Attachment')

    # Get the parent issue and check permissions
    issue = db.get_issue(attachment.get('issue_id'))
    if issue:
        user = get_current_user()
        audit_id = issue.get('audit_id') or get_active_audit_id()
        audit = db.get_audit(audit_id) if audit_id else None
        audit_context = {
            'id': audit_id,
            'auditor_id': audit.get('auditor_id') if audit else None,
            'reviewer_id': audit.get('reviewer_id') if audit else None
        }
        if not can_edit_record(user, audit_context, issue):
            return jsonify({'error': 'Cannot delete attachment - record is not editable'}), 403

    # Delete file from disk
    filepath = os.path.join(app.config['UPLOAD_FOLDER'], attachment['filename'])
    if os.path.exists(filepath):
        os.remove(filepath)

    # Delete from database
    db.delete_attachment(attachment_id)
    increment_data_version()
    return jsonify({'status': 'deleted'})

# ==================== Risk Attachments API ====================

@app.route('/api/risks/<risk_id>/attachments', methods=['GET'])
@require_login
def get_risk_attachments(risk_id):
    """Get all attachments for a risk."""
    attachments = db.get_attachments_for_risk(risk_id)
    return jsonify(attachments)

@app.route('/api/risks/<risk_id>/attachments', methods=['POST'])
@require_login
def upload_risk_attachment(risk_id):
    """Upload a file attachment for a risk. Enforces edit permissions."""
    risk = db.get_risk(risk_id)
    if not risk:
        return not_found_response('Risk')

    allowed, error = check_attachment_permission(risk, 'risk')
    if not allowed:
        return error

    if 'file' not in request.files:
        return jsonify({'error': 'No file provided'}), 400

    try:
        file_info = _process_file_upload(request.files['file'])
    except ValueError as e:
        return jsonify({'error': str(e)}), 400

    category = request.form.get('category', 'planning')
    if category not in ('planning', 'de', 'oe'):
        category = 'planning'

    attachment_id = db.add_risk_attachment(
        risk_id=risk_id,
        filename=file_info['unique_filename'],
        original_filename=file_info['original_filename'],
        file_size=file_info['file_size'],
        mime_type=file_info['mime_type'],
        description=request.form.get('description', ''),
        category=category,
        extracted_text=file_info['extracted_text']
    )

    increment_data_version()
    return jsonify(attachment_upload_response(file_info, attachment_id))

@app.route('/api/risk-attachments/<int:attachment_id>', methods=['GET'])
def download_risk_attachment(attachment_id):
    """Download a risk attachment file."""
    attachment = db.get_risk_attachment(attachment_id)
    if not attachment:
        return not_found_response('Attachment')

    return send_from_directory(
        app.config['UPLOAD_FOLDER'],
        attachment['filename'],
        download_name=attachment['original_filename'],
        as_attachment=True
    )

@app.route('/api/risk-attachments/<int:attachment_id>', methods=['DELETE'])
@require_login
def delete_risk_attachment(attachment_id):
    """Delete a risk attachment. Enforces edit permissions on the parent risk."""
    attachment = db.get_risk_attachment(attachment_id)
    if not attachment:
        return not_found_response('Attachment')

    # Get the parent risk and check permissions
    risk = db.get_risk(attachment.get('risk_id'))
    if risk:
        user = get_current_user()
        audit_id = risk.get('audit_id') or get_active_audit_id()
        audit = db.get_audit(audit_id) if audit_id else None
        audit_context = {
            'id': audit_id,
            'auditor_id': audit.get('auditor_id') if audit else None,
            'reviewer_id': audit.get('reviewer_id') if audit else None
        }
        if not can_edit_record(user, audit_context, risk):
            return jsonify({'error': 'Cannot delete attachment - record is not editable'}), 403

    # Delete file from disk
    filepath = os.path.join(app.config['UPLOAD_FOLDER'], attachment['filename'])
    if os.path.exists(filepath):
        os.remove(filepath)

    # Delete from database
    db.delete_risk_attachment(attachment_id)
    increment_data_version()
    return jsonify({'status': 'deleted'})

# ==================== Audit Attachments API ====================

@app.route('/api/audits/<int:audit_id>/attachments', methods=['GET'])
@require_login
def get_audit_attachments(audit_id):
    """Get all attachments for an audit."""
    attachments = db.get_attachments_for_audit(audit_id)
    return jsonify(attachments)

@app.route('/api/audits/<int:audit_id>/attachments', methods=['POST'])
@require_login
@require_non_viewer
def upload_audit_attachment(audit_id):
    """Upload a file attachment for an audit."""
    if 'file' not in request.files:
        return jsonify({'error': 'No file provided'}), 400

    try:
        file_info = _process_file_upload(request.files['file'])
    except ValueError as e:
        return jsonify({'error': str(e)}), 400

    attachment_id = db.add_audit_attachment(
        audit_id=audit_id,
        filename=file_info['unique_filename'],
        original_filename=file_info['original_filename'],
        file_size=file_info['file_size'],
        mime_type=file_info['mime_type'],
        description=request.form.get('description', ''),
        extracted_text=file_info['extracted_text']
    )

    increment_data_version()
    return jsonify(attachment_upload_response(file_info, attachment_id))

@app.route('/api/audit-attachments/<int:attachment_id>', methods=['GET'])
def download_audit_attachment(attachment_id):
    """Download an audit attachment file."""
    attachment = db.get_audit_attachment(attachment_id)
    if not attachment:
        return not_found_response('Attachment')

    return send_from_directory(
        app.config['UPLOAD_FOLDER'],
        attachment['filename'],
        download_name=attachment['original_filename'],
        as_attachment=True
    )

@app.route('/api/audit-attachments/<int:attachment_id>', methods=['DELETE'])
@require_login
@require_non_viewer
def delete_audit_attachment(attachment_id):
    """Delete an audit attachment."""
    attachment = db.get_audit_attachment(attachment_id)
    if not attachment:
        return not_found_response('Attachment')

    # Delete file from disk
    filepath = os.path.join(app.config['UPLOAD_FOLDER'], attachment['filename'])
    if os.path.exists(filepath):
        os.remove(filepath)

    # Delete from database
    db.delete_audit_attachment(attachment_id)
    increment_data_version()
    return jsonify({'status': 'deleted'})

@app.route('/api/issues/from-risk/<risk_id>', methods=['POST'])
def create_issue_from_risk(risk_id):
    """Create an issue from a RACM risk (used when checkbox is ticked)."""
    # Thread-safe data version update
    risk = db.get_risk(risk_id)
    if not risk:
        return jsonify({'error': 'Risk not found'}), 404

    # Check if issue already exists for this risk
    existing = db.get_issues_for_risk(risk_id)
    if existing:
        return jsonify({'status': 'exists', 'issue_id': existing[0]['issue_id']})

    # Create new issue with risk details
    issue_id = db.create_issue(
        risk_id=risk_id,
        title=f"Issue for {risk_id}: {risk.get('risk', 'No description')[:50]}",
        description=f"Risk: {risk.get('risk', '')}\nControl: {risk.get('control_id', '')}",
        severity='Medium',
        status='Open'
    )
    increment_data_version()
    return jsonify({'status': 'created', 'issue_id': issue_id})

# ==================== Export/Import ====================

@app.route('/api/export', methods=['GET'])
@require_login
def export_data():
    """Export all data as JSON."""
    return jsonify(db.export_all())

@app.route('/api/import', methods=['POST'])
@require_login
@require_non_viewer
def import_data():
    """Import data from JSON."""
    # Thread-safe data version update
    data = request.json
    clear = request.args.get('clear', 'false').lower() == 'true'
    db.import_all(data, clear_existing=clear)
    increment_data_version()
    return jsonify({'status': 'imported'})

# ==================== Chat API ====================

chat_history = []
_chat_history_lock = threading.Lock()


def add_to_chat_history(message: dict):
    """Add message to chat history with bounds (thread-safe)."""
    with _chat_history_lock:
        chat_history.append(message)
        # Trim to max size
        while len(chat_history) > MAX_CHAT_HISTORY:
            chat_history.pop(0)


def get_recent_chat_history(count: int = 10) -> list:
    """Get recent chat messages (thread-safe)."""
    with _chat_history_lock:
        return chat_history[-count:]

@app.route('/api/chat/status', methods=['GET'])
@require_login
def chat_status():
    """Check if API key is configured."""
    return jsonify({'configured': bool(ANTHROPIC_API_KEY)})

# ==================== Unified AI Tools & Functions ====================

def get_ai_tools():
    """Return the list of tools available to both Felix and sidebar chat."""
    return [
        # Data Creation Tools
        {
            "name": "add_racm_row",
            "description": "Add a new row to the Risk and Control Matrix (RACM).",
            "input_schema": {
                "type": "object",
                "properties": {
                    "risk_id": {"type": "string", "description": "Risk ID (e.g., R004)"},
                    "risk_description": {"type": "string", "description": "Description of the risk"},
                    "control_description": {"type": "string", "description": "Description of the control"},
                    "control_owner": {"type": "string", "description": "Who owns this control"},
                    "frequency": {"type": "string", "description": "How often (Daily, Weekly, Monthly, Quarterly)"},
                    "status": {"type": "string", "description": "Status (Effective, Needs Improvement, Ineffective, Not Tested)"}
                },
                "required": ["risk_id", "risk_description", "control_description"]
            }
        },
        {
            "name": "create_kanban_task",
            "description": "Create a new task on the Kanban board.",
            "input_schema": {
                "type": "object",
                "properties": {
                    "title": {"type": "string", "description": "Task title"},
                    "description": {"type": "string", "description": "Task description"},
                    "priority": {"type": "string", "enum": ["low", "medium", "high"]},
                    "assignee": {"type": "string", "description": "Who is assigned"},
                    "column": {"type": "string", "enum": ["planning", "fieldwork", "testing", "review", "complete"]},
                    "risk_id": {"type": "string", "description": "Link to RACM Risk ID (e.g., 'R001')"}
                },
                "required": ["title"]
            }
        },
        {
            "name": "create_flowchart",
            "description": "Create a new process flowchart.",
            "input_schema": {
                "type": "object",
                "properties": {
                    "name": {"type": "string", "description": "Flowchart name"},
                    "risk_id": {"type": "string", "description": "Link to RACM Risk ID"},
                    "steps": {
                        "type": "array",
                        "items": {
                            "type": "object",
                            "properties": {
                                "type": {"type": "string", "enum": ["start", "process", "decision", "control", "end"]},
                                "label": {"type": "string"},
                                "description": {"type": "string"}
                            },
                            "required": ["type", "label"]
                        }
                    }
                },
                "required": ["name", "steps"]
            }
        },
        {
            "name": "create_test_document",
            "description": "Create or update a test document (working paper) with testing procedures, findings, and evidence.",
            "input_schema": {
                "type": "object",
                "properties": {
                    "risk_code": {"type": "string", "description": "Risk ID (e.g., 'R001')"},
                    "doc_type": {"type": "string", "enum": ["de_testing", "oe_testing"], "description": "Type of testing document"},
                    "content": {"type": "string", "description": "The document content (can include HTML formatting)"}
                },
                "required": ["risk_code", "doc_type", "content"]
            }
        },
        {
            "name": "create_issue",
            "description": "Create a new issue in the issue log linked to a RACM risk.",
            "input_schema": {
                "type": "object",
                "properties": {
                    "risk_id": {"type": "string", "description": "RACM Risk ID to link to (e.g., 'R001')"},
                    "title": {"type": "string", "description": "Issue title"},
                    "description": {"type": "string", "description": "Issue description"},
                    "severity": {"type": "string", "enum": ["Low", "Medium", "High", "Critical"]},
                    "assigned_to": {"type": "string", "description": "Who should address this issue"},
                    "due_date": {"type": "string", "description": "Due date in YYYY-MM-DD format"}
                },
                "required": ["risk_id", "title"]
            }
        },
        # Data Update Tools
        {
            "name": "update_racm_status",
            "description": "Update the status of an existing RACM row.",
            "input_schema": {
                "type": "object",
                "properties": {
                    "risk_id": {"type": "string", "description": "Risk ID to update (e.g., R001)"},
                    "new_status": {"type": "string", "enum": ["Effective", "Needs Improvement", "Ineffective", "Not Tested"]}
                },
                "required": ["risk_id", "new_status"]
            }
        },
        {
            "name": "update_issue_documentation",
            "description": "Add or update documentation/evidence for an issue.",
            "input_schema": {
                "type": "object",
                "properties": {
                    "issue_id": {"type": "string", "description": "Issue ID (e.g., 'ISS-001')"},
                    "documentation": {"type": "string", "description": "The documentation content (can include HTML formatting)"}
                },
                "required": ["issue_id", "documentation"]
            }
        },
        {
            "name": "update_issue_status",
            "description": "Update the status of an issue.",
            "input_schema": {
                "type": "object",
                "properties": {
                    "issue_id": {"type": "string", "description": "Issue ID (e.g., 'ISS-001')"},
                    "new_status": {"type": "string", "enum": ["Open", "In Progress", "Resolved", "Closed"]}
                },
                "required": ["issue_id", "new_status"]
            }
        },
        {
            "name": "update_task",
            "description": "Update a Kanban task (move to different column, change priority, etc).",
            "input_schema": {
                "type": "object",
                "properties": {
                    "task_id": {"type": "integer", "description": "Task ID number"},
                    "title": {"type": "string", "description": "New title (optional)"},
                    "description": {"type": "string", "description": "New description (optional)"},
                    "priority": {"type": "string", "enum": ["low", "medium", "high"]},
                    "column": {"type": "string", "enum": ["planning", "fieldwork", "testing", "review", "complete"]},
                    "assignee": {"type": "string", "description": "New assignee (optional)"}
                },
                "required": ["task_id"]
            }
        },
        # Data Deletion Tools
        {
            "name": "delete_risk",
            "description": "Delete a risk from the RACM. WARNING: This also deletes associated tasks, issues, and attachments. Always confirm with the user before deleting.",
            "input_schema": {
                "type": "object",
                "properties": {
                    "risk_id": {"type": "string", "description": "Risk ID to delete (e.g., 'R001')"}
                },
                "required": ["risk_id"]
            }
        },
        {
            "name": "delete_task",
            "description": "Delete a task from the Kanban board.",
            "input_schema": {
                "type": "object",
                "properties": {
                    "task_id": {"type": "integer", "description": "Task ID to delete"}
                },
                "required": ["task_id"]
            }
        },
        {
            "name": "delete_issue",
            "description": "Delete an issue from the issue log. WARNING: This also deletes all attachments for the issue.",
            "input_schema": {
                "type": "object",
                "properties": {
                    "issue_id": {"type": "string", "description": "Issue ID to delete (e.g., 'ISS-001')"}
                },
                "required": ["issue_id"]
            }
        },
        {
            "name": "delete_flowchart",
            "description": "Delete a flowchart.",
            "input_schema": {
                "type": "object",
                "properties": {
                    "name": {"type": "string", "description": "Flowchart name to delete"}
                },
                "required": ["name"]
            }
        },
        {
            "name": "delete_attachment",
            "description": "Delete a file attachment from an issue or risk.",
            "input_schema": {
                "type": "object",
                "properties": {
                    "attachment_id": {"type": "integer", "description": "Attachment ID to delete"},
                    "attachment_type": {"type": "string", "enum": ["issue", "risk"], "description": "Type of attachment"}
                },
                "required": ["attachment_id", "attachment_type"]
            }
        },
        # Data Reading Tools
        {
            "name": "execute_sql",
            "description": "Execute a read-only SQL query to analyze audit data. Use for complex queries.",
            "input_schema": {
                "type": "object",
                "properties": {
                    "sql": {"type": "string", "description": "SELECT query to execute"}
                },
                "required": ["sql"]
            }
        },
        {
            "name": "get_audit_summary",
            "description": "Get a summary of current audit status.",
            "input_schema": {
                "type": "object",
                "properties": {},
                "required": []
            }
        },
        {
            "name": "read_test_document",
            "description": "Read the full content of a test document (working paper) for a specific risk. Use this to get DE or OE testing documentation before answering questions about testing.",
            "input_schema": {
                "type": "object",
                "properties": {
                    "risk_code": {"type": "string", "description": "Risk ID (e.g., 'R001')"},
                    "doc_type": {"type": "string", "enum": ["de_testing", "oe_testing"], "description": "Type of testing document"}
                },
                "required": ["risk_code", "doc_type"]
            }
        },
        {
            "name": "read_flowchart",
            "description": "Read the details of a flowchart including all nodes, steps, and descriptions. Use this to understand process flows before answering questions.",
            "input_schema": {
                "type": "object",
                "properties": {
                    "name": {"type": "string", "description": "Flowchart name"}
                },
                "required": ["name"]
            }
        },
        {
            "name": "read_issue_documentation",
            "description": "Read the documentation/evidence for an issue. Use this to get detailed findings, evidence, and supporting documentation for audit issues.",
            "input_schema": {
                "type": "object",
                "properties": {
                    "issue_id": {"type": "string", "description": "Issue ID (e.g., 'ISS-001')"}
                },
                "required": ["issue_id"]
            }
        },
        {
            "name": "list_issue_attachments",
            "description": "List all file attachments (evidence files like PDFs, Word docs, Excel, emails) for an issue.",
            "input_schema": {
                "type": "object",
                "properties": {
                    "issue_id": {"type": "string", "description": "Issue ID (e.g., 'ISS-001')"}
                },
                "required": ["issue_id"]
            }
        },
        {
            "name": "list_risk_attachments",
            "description": "List all file attachments (evidence files like PDFs, Word docs, Excel, emails) for a risk in the RACM.",
            "input_schema": {
                "type": "object",
                "properties": {
                    "risk_id": {"type": "string", "description": "Risk ID (e.g., 'R001')"}
                },
                "required": ["risk_id"]
            }
        },
        {
            "name": "read_attachment_content",
            "description": "Read the extracted text content from an uploaded attachment file. Use this to analyze document contents like PDFs, Word documents, or Excel spreadsheets that have been uploaded as evidence.",
            "input_schema": {
                "type": "object",
                "properties": {
                    "attachment_id": {"type": "integer", "description": "The attachment ID number"},
                    "attachment_type": {"type": "string", "enum": ["issue", "risk"], "description": "Whether this is an issue attachment or risk attachment"}
                },
                "required": ["attachment_id", "attachment_type"]
            }
        },
        # Audit Library Tool (RAG Search)
        {
            "name": "search_audit_library",
            "description": "Search the Audit Library for relevant information from reference documents like COBIT, COSO, IIA Standards, audit methodologies, policies, and other uploaded frameworks. Use this to find guidance, best practices, requirements, or standards that can inform audit decisions, testing approaches, or recommendations. The library uses semantic search to find the most relevant content.",
            "input_schema": {
                "type": "object",
                "properties": {
                    "query": {"type": "string", "description": "The search query - describe what information you're looking for (e.g., 'IT general controls testing requirements', 'segregation of duties best practices', 'audit evidence standards')"},
                    "limit": {"type": "integer", "description": "Maximum number of results to return (default: 5, max: 10)"}
                },
                "required": ["query"]
            }
        },
        # User Interaction Tool
        {
            "name": "ask_clarifying_question",
            "description": "Ask the user a clarifying question when you need more information to complete their request accurately. Use this when the request is ambiguous, you need to confirm which item to act on, or before performing destructive actions like deletes.",
            "input_schema": {
                "type": "object",
                "properties": {
                    "question": {"type": "string", "description": "The question to ask the user"},
                    "options": {
                        "type": "array",
                        "items": {"type": "string"},
                        "description": "Optional list of choices to present to the user"
                    }
                },
                "required": ["question"]
            }
        }
    ]


def build_ai_system_prompt(context, include_capabilities=True):
    """Build the system prompt for AI with audit context summary.

    Note: Only summaries are included to reduce token usage.
    Felix should use tools (execute_sql, get_audit_summary) to fetch detailed data when needed.
    """
    # Build concise risk list (just IDs and brief info)
    risk_list = []
    for r in context['risks'][:20]:  # Limit to 20 risks
        risk_list.append(f"- {r.get('risk_id', 'N/A')}: {r.get('risk', 'No description')[:60]}... [{r.get('status', 'Unknown')}]")
    risk_summary_text = "\n".join(risk_list) if risk_list else "No risks in RACM"
    if len(context['risks']) > 20:
        risk_summary_text += f"\n... and {len(context['risks']) - 20} more risks (use execute_sql to see all)"

    # Build audit plan list
    audit_summary = context.get('audit_summary', {})
    audit_list = []
    for a in context.get('audits', []):
        audit_list.append(f"- **{a.get('title', 'Untitled')}** [{a.get('quarter', 'N/A')}] - {a.get('status', 'unknown')} | {a.get('owner', 'Unassigned')} | {a.get('audit_area', 'N/A')}")
    audit_plan_text = "\n".join(audit_list) if audit_list else "No audits in the annual plan"

    prompt = f"""You are Felix, an intelligent AI audit assistant with FULL ACCESS to the SmartPapers audit database.

## Current Audit Summary:
- **Annual Audit Plan:** {audit_summary.get('total', 0)} audits - By Quarter: {json.dumps(audit_summary.get('by_quarter', {}))} | Status: {json.dumps(audit_summary.get('by_status', {}))}
- **Risks:** {context['risk_summary']['total']} total - {json.dumps(context['risk_summary']['by_status'])}
- **Tasks:** {context['task_summary']['total']} total - {json.dumps(context['task_summary']['by_column'])}
- **Issues:** {context['issue_summary']['total']} total - {json.dumps(context['issue_summary']['by_status'])}
- **Flowcharts:** {context['flowchart_count']}
- **Test Documents:** {context['test_doc_count']}
- **Attachments:** {context['issue_attachment_count']} issue files, {context['risk_attachment_count']} risk files

## Annual Audit Plan (All Planned Audits):
{audit_plan_text}

## RACM Overview (use execute_sql for full details):
{risk_summary_text}

## Available Tools for Data Access:
- **execute_sql**: Query the database for detailed risk, task, or issue data
- **get_audit_summary**: Get current statistics
- **read_test_document / read_flowchart**: Read specific documents
- **search_audit_library**: Search reference documents (COBIT, COSO, IIA Standards)
"""

    if include_capabilities:
        prompt += """
## Your Capabilities - USE TOOLS TO TAKE ACTION:
You have FULL access to create, read, update, and delete audit data:

**CREATE:** Add risks, tasks, issues, flowcharts, test documents
**READ:** Query data with SQL, read documents, flowcharts, attachments
**UPDATE:** Change statuses, update documentation, move tasks
**DELETE:** Remove risks, tasks, issues, flowcharts, attachments
**AUDIT LIBRARY:** Search reference documents (COBIT, COSO, IIA Standards, methodologies) for guidance

## IMPORTANT GUIDELINES:

1. **Take Action:** When the user asks you to do something, USE THE TOOLS to do it. Don't just describe what you would do.

2. **Ask When Unclear:** If a request is ambiguous (e.g., "update the status" without specifying which risk), use ask_clarifying_question to get details.

3. **Confirm Destructive Actions:** Before deleting anything, use ask_clarifying_question to confirm with the user.

4. **Seamless Document Access:** When asked about testing or documentation:
   - Check if documents exist in the metadata above
   - Use read_test_document, read_flowchart, or read_issue_documentation to fetch content BEFORE answering
   - The user should get complete answers without asking you to fetch documents

5. **Be Proactive:** Offer to create missing documents, suggest next steps, and help streamline the audit process.

6. **Use the Audit Library:** When the user asks about audit standards, best practices, control testing approaches, or needs guidance on methodology, use search_audit_library to find relevant content from reference documents before answering.

Be concise and professional. Use markdown formatting for clarity."""

    return prompt


@app.route('/api/chat', methods=['POST'])
@require_login
@require_non_viewer
def chat():
    # Thread-safe data version update
    data = request.json
    user_message = data.get('message', '')
    api_key = ANTHROPIC_API_KEY

    if not api_key:
        return jsonify({'error': 'No API key provided. Set ANTHROPIC_API_KEY environment variable.'})

    # Get full context and use shared functions
    context = db.get_full_context()
    system_prompt = build_ai_system_prompt(context)
    tools = get_ai_tools()

    try:
        client = anthropic.Anthropic(api_key=api_key)
        add_to_chat_history({"role": "user", "content": user_message})
        messages = get_recent_chat_history(10)

        response = client.messages.create(
            model=CLAUDE_MODEL,
            max_tokens=2048,
            system=system_prompt,
            tools=tools,
            messages=messages
        )

        final_response = ""
        tool_results = []

        while response.stop_reason == "tool_use":
            for content in response.content:
                if content.type == "tool_use":
                    result = execute_tool(content.name, content.input)
                    tool_results.append({
                        "type": "tool_result",
                        "tool_use_id": content.id,
                        "content": result
                    })
                elif content.type == "text":
                    final_response += content.text

            messages.append({"role": "assistant", "content": response.content})
            messages.append({"role": "user", "content": tool_results})

            response = client.messages.create(
                model=CLAUDE_MODEL,
                max_tokens=2048,
                system=system_prompt,
                tools=tools,
                messages=messages
            )
            tool_results = []

        for content in response.content:
            if hasattr(content, 'text'):
                final_response += content.text

        add_to_chat_history({"role": "assistant", "content": final_response})
        return jsonify({'response': final_response, 'data_version': get_data_version()})

    except anthropic.AuthenticationError:
        return jsonify({'error': 'AI service authentication failed. Please check API key configuration.'}), 401
    except anthropic.RateLimitError:
        return jsonify({'error': 'AI service rate limit exceeded. Please try again in a moment.'}), 429
    except anthropic.APIError as e:
        app.logger.error(f"Anthropic API error: {e}")
        return jsonify({'error': 'AI service temporarily unavailable.'}), 503
    except Exception as e:
        app.logger.error(f"Chat error: {e}")
        return jsonify({'error': 'An error occurred processing your request.'}), 500


def execute_tool(tool_name, tool_input):
    """Execute a tool and return the result."""
    # Thread-safe data version update

    if tool_name == "add_racm_row":
        db.create_risk(
            risk_id=tool_input.get('risk_id', ''),
            risk=tool_input.get('risk_description', ''),
            control_id=tool_input.get('control_description', ''),
            control_owner=tool_input.get('control_owner', ''),
            status=tool_input.get('status', 'Not Tested')
        )
        increment_data_version()
        return f"Successfully added RACM row: {tool_input.get('risk_id')} - {tool_input.get('risk_description')}"

    elif tool_name == "create_kanban_task":
        risk_id = tool_input.get('risk_id', '')
        task_id = db.create_task(
            title=tool_input.get('title', 'New Task'),
            description=tool_input.get('description', ''),
            priority=tool_input.get('priority', 'medium'),
            assignee=tool_input.get('assignee', ''),
            column_id=tool_input.get('column', 'planning'),
            risk_id=risk_id if risk_id else None
        )
        increment_data_version()
        linked_msg = f" Linked to RACM row {risk_id}." if risk_id else ""
        return f"Successfully created task '{tool_input.get('title')}' (ID: {task_id}).{linked_msg}"

    elif tool_name == "update_racm_status":
        risk_id = tool_input.get('risk_id', '').upper()
        new_status = tool_input.get('new_status', '')
        if db.update_risk(risk_id, status=new_status):
            increment_data_version()
            return f"Successfully updated {risk_id} status to '{new_status}'"
        return f"Risk ID '{risk_id}' not found"

    elif tool_name == "execute_sql":
        sql = tool_input.get('sql', '')
        try:
            results = db.execute_query(sql)
            return json.dumps(results, indent=2)
        except Exception as e:
            return f"Query error: {str(e)}"

    elif tool_name == "get_audit_summary":
        summary = {
            'risks': db.get_risk_summary(),
            'tasks': db.get_task_summary(),
            'flowcharts': len(db.get_all_flowcharts())
        }
        return json.dumps(summary, indent=2)

    elif tool_name == "create_flowchart":
        name = tool_input.get('name', 'new-flowchart').lower().replace(' ', '-')
        steps = tool_input.get('steps', [])
        risk_id = tool_input.get('risk_id', '')

        # Build Drawflow data structure
        drawflow_data = {"drawflow": {"Home": {"data": {}}}}
        nodes = drawflow_data["drawflow"]["Home"]["data"]

        node_config = {
            'start': {'inputs': 0, 'outputs': 1},
            'process': {'inputs': 1, 'outputs': 1},
            'decision': {'inputs': 1, 'outputs': 2},
            'control': {'inputs': 1, 'outputs': 1},
            'end': {'inputs': 1, 'outputs': 0},
        }

        y_pos = 50
        prev_node_id = None

        for i, step in enumerate(steps):
            node_id = str(i + 1)
            node_type = step.get('type', 'process')
            label = step.get('label', 'Step')
            description = step.get('description', '')

            config = node_config.get(node_type, {'inputs': 1, 'outputs': 1})

            inputs = {}
            if config['inputs'] > 0:
                inputs["input_1"] = {"connections": []}
                if prev_node_id:
                    inputs["input_1"]["connections"].append({
                        "node": prev_node_id,
                        "input": "output_1"
                    })

            outputs = {}
            for j in range(config['outputs']):
                outputs[f"output_{j+1}"] = {"connections": []}

            nodes[node_id] = {
                "id": int(node_id),
                "name": node_type,
                "data": {"name": label, "description": description},
                "class": node_type,
                "html": f'<div class="node-header"><input type="text" class="header-input" value="{label}" df-name></div><div class="node-body"><textarea class="node-textarea" placeholder="Description..." df-description>{description}</textarea></div>',
                "typenode": False,
                "inputs": inputs,
                "outputs": outputs,
                "pos_x": 150,
                "pos_y": y_pos
            }

            if prev_node_id and prev_node_id in nodes:
                nodes[prev_node_id]["outputs"]["output_1"]["connections"].append({
                    "node": node_id,
                    "output": "input_1"
                })

            prev_node_id = node_id
            y_pos += 150

        db.save_flowchart(name, drawflow_data, risk_id if risk_id else None)
        increment_data_version()

        linked_msg = f" Linked to RACM row {risk_id}." if risk_id else ""
        return f"Successfully created flowchart '{name}' with {len(steps)} nodes.{linked_msg} View at: /flowchart/{name}"

    elif tool_name == "read_test_document":
        risk_code = tool_input.get('risk_code', '').upper()
        doc_type = tool_input.get('doc_type', '')

        if doc_type not in ('de_testing', 'oe_testing'):
            return f"Invalid doc_type '{doc_type}'. Must be 'de_testing' or 'oe_testing'."

        doc = db.get_test_document_by_risk_code(risk_code, doc_type)
        if not doc:
            return f"No {doc_type} document found for risk {risk_code}."

        content = doc.get('content', '')
        if not content.strip():
            return f"The {doc_type} document for {risk_code} exists but is empty."

        # Strip HTML tags for cleaner AI reading
        text_content = re.sub(r'<[^>]+>', '', content)
        text_content = re.sub(r'\s+', ' ', text_content).strip()

        doc_type_label = "Design Effectiveness Testing" if doc_type == "de_testing" else "Operational Effectiveness Testing"
        return f"## {doc_type_label} Document for {risk_code}\n\n{text_content}"

    elif tool_name == "create_test_document":
        risk_code = tool_input.get('risk_code', '').upper()
        doc_type = tool_input.get('doc_type', '')
        content = tool_input.get('content', '')

        if doc_type not in ('de_testing', 'oe_testing'):
            return f"Invalid doc_type '{doc_type}'. Must be 'de_testing' or 'oe_testing'."

        if not content.strip():
            return "Cannot create an empty document. Please provide content."

        # Wrap plain text in HTML paragraph if needed
        if not content.strip().startswith('<'):
            content = f"<p>{content}</p>"

        doc_id = db.save_test_document_by_risk_code(risk_code, doc_type, content)
        if doc_id is None:
            return f"Risk {risk_code} not found. Cannot create document."

        increment_data_version()
        doc_type_label = "Design Effectiveness Testing" if doc_type == "de_testing" else "Operational Effectiveness Testing"
        return f"Successfully created/updated {doc_type_label} document for {risk_code}."

    elif tool_name == "read_flowchart":
        name = tool_input.get('name', '')
        fc = db.get_flowchart_with_details(name)

        if not fc:
            return f"Flowchart '{name}' not found."

        nodes = fc.get('nodes', [])
        if not nodes:
            return f"Flowchart '{name}' exists but has no nodes."

        result = f"## Flowchart: {name}\n"
        if fc.get('risk_id'):
            result += f"Linked to Risk: {fc['risk_id']}\n"
        result += f"\n### Process Steps ({len(nodes)} nodes):\n"

        for node in nodes:
            node_type = node.get('type', 'unknown')
            label = node.get('label', 'Unnamed')
            desc = node.get('description', '')
            result += f"\n- **[{node_type.upper()}]** {label}"
            if desc:
                result += f"\n  {desc}"

        return result

    elif tool_name == "read_issue_documentation":
        issue_id = tool_input.get('issue_id', '').upper()
        issue = db.get_issue(issue_id)

        if not issue:
            return f"Issue '{issue_id}' not found."

        doc = issue.get('documentation', '')
        if not doc or not doc.strip():
            return f"Issue {issue_id} exists but has no documentation yet."

        # Strip HTML tags for cleaner AI reading
        text_content = re.sub(r'<[^>]+>', '', doc)
        text_content = re.sub(r'\s+', ' ', text_content).strip()

        return f"## Documentation for Issue {issue_id}\n\nTitle: {issue.get('title', 'N/A')}\nRisk: {issue.get('risk_id', 'N/A')}\nSeverity: {issue.get('severity', 'N/A')}\nStatus: {issue.get('status', 'N/A')}\n\n### Evidence/Findings:\n{text_content}"

    elif tool_name == "update_issue_documentation":
        issue_id = tool_input.get('issue_id', '').upper()
        documentation = tool_input.get('documentation', '')

        if not documentation.strip():
            return "Cannot save empty documentation. Please provide content."

        # Wrap plain text in HTML paragraph if needed
        if not documentation.strip().startswith('<'):
            documentation = f"<p>{documentation}</p>"

        if db.save_issue_documentation(issue_id, documentation):
            increment_data_version()
            return f"Successfully saved documentation for issue {issue_id}."
        return f"Issue {issue_id} not found."

    elif tool_name == "list_issue_attachments":
        issue_id = tool_input.get('issue_id', '').upper()
        attachments = db.get_attachments_for_issue(issue_id)
        return _format_attachment_list(attachments, 'issue', issue_id)

    elif tool_name == "list_risk_attachments":
        risk_id = tool_input.get('risk_id', '').upper()
        attachments = db.get_attachments_for_risk(risk_id)
        return _format_attachment_list(attachments, 'risk', risk_id)

    elif tool_name == "read_attachment_content":
        attachment_id = tool_input.get('attachment_id')
        attachment_type = tool_input.get('attachment_type', 'issue')

        if attachment_type == 'issue':
            attachment = db.get_attachment(attachment_id)
        else:
            attachment = db.get_risk_attachment(attachment_id)

        if not attachment:
            return f"Attachment with ID {attachment_id} not found."

        filename = attachment.get('original_filename', 'Unknown file')
        extracted_text = attachment.get('extracted_text', '')

        if not extracted_text:
            return f"No text content available for '{filename}'. The file may be an image or unsupported format."

        if extracted_text.startswith('['):
            # This is an error/status message from extraction
            return f"Cannot read content of '{filename}': {extracted_text}"

        # Truncate very long content for the AI
        if len(extracted_text) > 30000:
            extracted_text = extracted_text[:30000] + "\n\n[Content truncated - file contains more text]"

        return f"## Content of '{filename}'\n\n{extracted_text}"

    # New tools for unified AI
    elif tool_name == "create_issue":
        risk_id = tool_input.get('risk_id', '').upper()
        title = tool_input.get('title', '')
        description = tool_input.get('description', '')
        severity = tool_input.get('severity', 'Medium')
        assigned_to = tool_input.get('assigned_to', '')
        due_date = tool_input.get('due_date', '')

        # Verify risk exists
        risk = db.get_risk(risk_id)
        if not risk:
            return f"Risk {risk_id} not found. Cannot create issue."

        # Generate issue ID
        with db._connection() as conn:
            cursor = conn.execute("SELECT COUNT(*) FROM issues")
            count = cursor.fetchone()[0] + 1
            issue_id = f"ISS-{count:03d}"

            conn.execute('''
                INSERT INTO issues (issue_id, risk_id, title, description, severity, status, assigned_to, due_date)
                VALUES (?, ?, ?, ?, ?, 'Open', ?, ?)
            ''', (issue_id, risk_id, title, description, severity, assigned_to, due_date if due_date else None))
            conn.commit()

        increment_data_version()
        return f"Successfully created issue {issue_id}: '{title}' linked to {risk_id}."

    elif tool_name == "update_issue_status":
        issue_id = tool_input.get('issue_id', '').upper()
        new_status = tool_input.get('new_status', '')

        with db._connection() as conn:
            cursor = conn.execute("UPDATE issues SET status = ?, updated_at = CURRENT_TIMESTAMP WHERE issue_id = ?", (new_status, issue_id))
            conn.commit()
            if cursor.rowcount > 0:
                increment_data_version()
                return f"Successfully updated {issue_id} status to '{new_status}'."
        return f"Issue {issue_id} not found."

    elif tool_name == "update_task":
        task_id = tool_input.get('task_id')
        updates = {}
        if 'title' in tool_input:
            updates['title'] = tool_input['title']
        if 'description' in tool_input:
            updates['description'] = tool_input['description']
        if 'priority' in tool_input:
            updates['priority'] = tool_input['priority']
        if 'column' in tool_input:
            updates['column_id'] = tool_input['column']
        if 'assignee' in tool_input:
            updates['assignee'] = tool_input['assignee']

        if not updates:
            return "No updates provided for task."

        set_clause = ", ".join([f"{k} = ?" for k in updates.keys()])
        values = list(updates.values()) + [task_id]

        with db._connection() as conn:
            cursor = conn.execute(f"UPDATE tasks SET {set_clause}, updated_at = CURRENT_TIMESTAMP WHERE id = ?", values)
            conn.commit()
            if cursor.rowcount > 0:
                increment_data_version()
                return f"Successfully updated task {task_id}."
        return f"Task {task_id} not found."

    elif tool_name == "delete_risk":
        risk_id = tool_input.get('risk_id', '').upper()
        with db._connection() as conn:
            # Check if exists
            cursor = conn.execute("SELECT id FROM risks WHERE risk_id = ?", (risk_id,))
            if not cursor.fetchone():
                return f"Risk {risk_id} not found."

            # Delete associated data
            conn.execute("DELETE FROM tasks WHERE risk_id = (SELECT id FROM risks WHERE risk_id = ?)", (risk_id,))
            conn.execute("DELETE FROM issues WHERE risk_id = ?", (risk_id,))
            conn.execute("DELETE FROM risk_attachments WHERE risk_id = ?", (risk_id,))
            conn.execute("DELETE FROM test_documents WHERE risk_id = (SELECT id FROM risks WHERE risk_id = ?)", (risk_id,))
            conn.execute("DELETE FROM flowcharts WHERE risk_id = (SELECT id FROM risks WHERE risk_id = ?)", (risk_id,))
            conn.execute("DELETE FROM risks WHERE risk_id = ?", (risk_id,))
            conn.commit()

        increment_data_version()
        return f"Successfully deleted risk {risk_id} and all associated data."

    elif tool_name == "delete_task":
        task_id = tool_input.get('task_id')
        with db._connection() as conn:
            cursor = conn.execute("DELETE FROM tasks WHERE id = ?", (task_id,))
            conn.commit()
            if cursor.rowcount > 0:
                increment_data_version()
                return f"Successfully deleted task {task_id}."
        return f"Task {task_id} not found."

    elif tool_name == "delete_issue":
        issue_id = tool_input.get('issue_id', '').upper()
        with db._connection() as conn:
            # Delete attachments first
            conn.execute("DELETE FROM issue_attachments WHERE issue_id = ?", (issue_id,))
            cursor = conn.execute("DELETE FROM issues WHERE issue_id = ?", (issue_id,))
            conn.commit()
            if cursor.rowcount > 0:
                increment_data_version()
                return f"Successfully deleted issue {issue_id} and its attachments."
        return f"Issue {issue_id} not found."

    elif tool_name == "delete_flowchart":
        name = tool_input.get('name', '').lower().replace(' ', '-')
        with db._connection() as conn:
            cursor = conn.execute("DELETE FROM flowcharts WHERE name = ?", (name,))
            conn.commit()
            if cursor.rowcount > 0:
                increment_data_version()
                return f"Successfully deleted flowchart '{name}'."
        return f"Flowchart '{name}' not found."

    elif tool_name == "delete_attachment":
        attachment_id = tool_input.get('attachment_id')
        attachment_type = tool_input.get('attachment_type', 'issue')

        table = 'issue_attachments' if attachment_type == 'issue' else 'risk_attachments'
        with db._connection() as conn:
            cursor = conn.execute(f"DELETE FROM {table} WHERE id = ?", (attachment_id,))
            conn.commit()
            if cursor.rowcount > 0:
                increment_data_version()
                return f"Successfully deleted {attachment_type} attachment {attachment_id}."
        return f"Attachment {attachment_id} not found."

    elif tool_name == "search_audit_library":
        query = tool_input.get('query', '')
        limit = min(tool_input.get('limit', 5), 10)  # Cap at 10

        if not query.strip():
            return "Please provide a search query."

        try:
            # Try semantic search first (requires embeddings)
            query_embedding = generate_embedding(query)
            results = db.search_library(query_embedding, limit=limit)

            if not results:
                # Fall back to keyword search
                results = db.search_library_keyword(query, limit=limit)

            if not results:
                return f"No results found in the Audit Library for: '{query}'. Try different keywords or ensure documents have been uploaded to the library."

            # Format results for the AI
            response = f"## Audit Library Search Results for: '{query}'\n\n"
            for i, result in enumerate(results, 1):
                response += f"### Result {i}: {result.get('document_name', 'Unknown Document')}\n"
                if result.get('section'):
                    response += f"**Section:** {result['section']}\n"
                response += f"**Source:** {result.get('source', 'Unknown')}\n"
                response += f"**Type:** {result.get('doc_type', 'Unknown')}\n"
                if result.get('similarity'):
                    response += f"**Relevance:** {result['similarity']*100:.0f}%\n"
                response += f"\n{result.get('content', '')}\n\n---\n\n"

            return response

        except Exception as e:
            # If embedding fails, try keyword search
            try:
                results = db.search_library_keyword(query, limit=limit)
                if results:
                    response = f"## Audit Library Search Results (keyword) for: '{query}'\n\n"
                    for i, result in enumerate(results, 1):
                        response += f"### Result {i}: {result.get('document_name', 'Unknown Document')}\n"
                        if result.get('section'):
                            response += f"**Section:** {result['section']}\n"
                        response += f"\n{result.get('content', '')}\n\n---\n\n"
                    return response
                return f"No results found in the Audit Library for: '{query}'."
            except Exception as e2:
                return f"Library search error: {str(e2)}"

    elif tool_name == "ask_clarifying_question":
        # This tool returns a special marker that the frontend handles
        question = tool_input.get('question', '')
        options = tool_input.get('options', [])
        return json.dumps({
            "type": "clarifying_question",
            "question": question,
            "options": options
        })

    return "Unknown tool"


@app.route('/api/chat/clear', methods=['POST'])
@require_login
def clear_chat():
    global chat_history
    chat_history = []
    return jsonify({'status': 'cleared'})


# ==================== Felix AI Chat ====================

def init_felix_tables():
    """Initialize Felix AI chat tables."""
    with db._connection() as conn:
        conn.executescript('''
            CREATE TABLE IF NOT EXISTS felix_conversations (
                id TEXT PRIMARY KEY,
                user_id TEXT NOT NULL DEFAULT 'default_user',
                title TEXT DEFAULT 'New Chat',
                created_at DATETIME DEFAULT CURRENT_TIMESTAMP,
                updated_at DATETIME DEFAULT CURRENT_TIMESTAMP
            );

            CREATE TABLE IF NOT EXISTS felix_messages (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                conversation_id TEXT NOT NULL,
                role TEXT NOT NULL,
                content TEXT NOT NULL,
                timestamp DATETIME DEFAULT CURRENT_TIMESTAMP,
                FOREIGN KEY (conversation_id) REFERENCES felix_conversations(id) ON DELETE CASCADE
            );

            CREATE INDEX IF NOT EXISTS idx_felix_messages_conv ON felix_messages(conversation_id);
            CREATE INDEX IF NOT EXISTS idx_felix_conversations_user ON felix_conversations(user_id);

            CREATE TABLE IF NOT EXISTS felix_attachments (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                conversation_id TEXT NOT NULL,
                filename TEXT NOT NULL,
                original_filename TEXT NOT NULL,
                file_size INTEGER,
                mime_type TEXT,
                extracted_text TEXT,
                uploaded_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                FOREIGN KEY (conversation_id) REFERENCES felix_conversations(id) ON DELETE CASCADE
            );

            CREATE INDEX IF NOT EXISTS idx_felix_attachments_conv ON felix_attachments(conversation_id);
        ''')

# Initialize Felix tables
init_felix_tables()


@app.route('/felix')
@require_login
def felix_chat():
    """Felix AI full-screen chat page."""
    user = get_current_user()
    user_id = user['id'] if user else 'default_user'
    ctx = get_audit_context()
    return render_template('felix.html', user_id=user_id, active_page='felix', **ctx)


@app.route('/api/felix/conversations', methods=['GET'])
@require_login
@require_non_viewer
def get_felix_conversations():
    """Get all conversations for current user."""
    user_id = session.get('user_id', 'default_user')
    with db._connection() as conn:
        cursor = conn.execute('''
            SELECT id, title, updated_at FROM felix_conversations
            WHERE user_id = ? ORDER BY updated_at DESC LIMIT 50
        ''', (user_id,))
        return jsonify([dict(row) for row in cursor.fetchall()])


@app.route('/api/felix/conversations', methods=['POST'])
@require_login
@require_non_viewer
def create_felix_conversation():
    """Create a new conversation."""
    user_id = session.get('user_id', 'default_user')
    conv_id = str(uuid.uuid4())
    with db._connection() as conn:
        conn.execute(
            'INSERT INTO felix_conversations (id, user_id, title) VALUES (?, ?, ?)',
            (conv_id, user_id, 'New Chat')
        )
    return jsonify({'id': conv_id, 'title': 'New Chat'})


@app.route('/api/felix/conversations/<conv_id>', methods=['DELETE'])
@require_login
@require_non_viewer
def delete_felix_conversation(conv_id):
    """Delete a conversation and its messages."""
    with db._connection() as conn:
        conn.execute('DELETE FROM felix_messages WHERE conversation_id = ?', (conv_id,))
        conn.execute('DELETE FROM felix_conversations WHERE id = ?', (conv_id,))
    return jsonify({'success': True})


@app.route('/api/felix/conversations/<conv_id>/messages', methods=['GET'])
@require_login
@require_non_viewer
def get_felix_messages(conv_id):
    """Get all messages for a conversation."""
    with db._connection() as conn:
        cursor = conn.execute('''
            SELECT role, content, timestamp FROM felix_messages
            WHERE conversation_id = ? ORDER BY timestamp ASC
        ''', (conv_id,))
        return jsonify([dict(row) for row in cursor.fetchall()])


@app.route('/api/felix/conversations/<conv_id>/messages', methods=['POST'])
@require_login
@require_non_viewer
def send_felix_message(conv_id):
    """Send a message and get AI response."""
    data = request.json
    content = data.get('content', '').strip()
    if not content:
        return jsonify({'error': 'Empty message'}), 400

    # Save user message and get history + attachments
    with db._connection() as conn:
        conn.execute(
            'INSERT INTO felix_messages (conversation_id, role, content) VALUES (?, ?, ?)',
            (conv_id, 'user', content)
        )
        cursor = conn.execute('''
            SELECT role, content FROM felix_messages
            WHERE conversation_id = ? ORDER BY timestamp ASC LIMIT 20
        ''', (conv_id,))
        history = [dict(row) for row in cursor.fetchall()]

        # Fetch conversation attachments
        cursor = conn.execute('''
            SELECT original_filename, extracted_text FROM felix_attachments
            WHERE conversation_id = ? ORDER BY uploaded_at ASC
        ''', (conv_id,))
        attachments = [dict(row) for row in cursor.fetchall()]

    # Call Claude API with attachments context
    try:
        response_text = call_felix_ai(history, attachments)
    except anthropic.AuthenticationError:
        return jsonify({'error': 'AI service authentication failed.'}), 401
    except anthropic.RateLimitError:
        return jsonify({'error': 'AI service rate limit exceeded. Please try again.'}), 429
    except anthropic.APIError as e:
        app.logger.error(f"Felix API error: {e}")
        return jsonify({'error': 'AI service temporarily unavailable.'}), 503
    except Exception as e:
        app.logger.error(f"Felix chat error: {e}")
        return jsonify({'error': 'An error occurred processing your message.'}), 500

    # Save assistant response and update conversation
    with db._connection() as conn:
        conn.execute(
            'INSERT INTO felix_messages (conversation_id, role, content) VALUES (?, ?, ?)',
            (conv_id, 'assistant', response_text)
        )
        # Update title if first message
        if len(history) <= 1:
            title = content[:40] + ('...' if len(content) > 40 else '')
            conn.execute(
                'UPDATE felix_conversations SET title = ?, updated_at = CURRENT_TIMESTAMP WHERE id = ?',
                (title, conv_id)
            )
        else:
            conn.execute(
                'UPDATE felix_conversations SET updated_at = CURRENT_TIMESTAMP WHERE id = ?',
                (conv_id,)
            )

    return jsonify({'role': 'assistant', 'content': response_text, 'data_version': get_data_version()})


# ===== Felix Attachment Routes =====

@app.route('/api/felix/conversations/<conv_id>/attachments', methods=['GET'])
@require_login
def get_felix_attachments(conv_id):
    """Get all attachments for a conversation."""
    with db._connection() as conn:
        cursor = conn.execute('''
            SELECT id, original_filename, file_size, mime_type, uploaded_at
            FROM felix_attachments WHERE conversation_id = ?
            ORDER BY uploaded_at DESC
        ''', (conv_id,))
        return jsonify([dict(row) for row in cursor.fetchall()])


@app.route('/api/felix/conversations/<conv_id>/attachments', methods=['POST'])
@require_login
@require_non_viewer
def upload_felix_attachment(conv_id):
    """Upload file(s) to a conversation."""
    if 'file' not in request.files:
        return jsonify({'error': 'No file provided'}), 400

    file = request.files['file']
    if file.filename == '':
        return jsonify({'error': 'No file selected'}), 400

    # Use existing file processing
    try:
        file_info = _process_file_upload(file)
    except ValueError as e:
        return jsonify({'error': str(e)}), 400

    # Save to database
    with db._connection() as conn:
        cursor = conn.execute('''
            INSERT INTO felix_attachments
            (conversation_id, filename, original_filename, file_size, mime_type, extracted_text)
            VALUES (?, ?, ?, ?, ?, ?)
        ''', (
            conv_id,
            file_info['unique_filename'],
            file_info['original_filename'],
            file_info['file_size'],
            file_info['mime_type'],
            file_info.get('extracted_text', '')
        ))
        attachment_id = cursor.lastrowid

    return jsonify({
        'id': attachment_id,
        'original_filename': file_info['original_filename'],
        'file_size': file_info['file_size'],
        'mime_type': file_info['mime_type']
    })


@app.route('/api/felix/attachments/<int:attachment_id>', methods=['DELETE'])
@require_login
@require_non_viewer
def delete_felix_attachment(attachment_id):
    """Delete an attachment."""
    with db._connection() as conn:
        # Get filename to delete from disk
        cursor = conn.execute(
            'SELECT filename FROM felix_attachments WHERE id = ?',
            (attachment_id,)
        )
        row = cursor.fetchone()
        if not row:
            return jsonify({'error': 'Attachment not found'}), 404

        # Delete file from disk
        filepath = os.path.join(app.config['UPLOAD_FOLDER'], row['filename'])
        if os.path.exists(filepath):
            os.remove(filepath)

        # Delete from database
        conn.execute('DELETE FROM felix_attachments WHERE id = ?', (attachment_id,))

    return jsonify({'success': True})


@app.route('/api/felix/attachments/<int:attachment_id>/download')
@require_login
def download_felix_attachment(attachment_id):
    """Download an attachment."""
    with db._connection() as conn:
        cursor = conn.execute(
            'SELECT filename, original_filename, mime_type FROM felix_attachments WHERE id = ?',
            (attachment_id,)
        )
        row = cursor.fetchone()
        if not row:
            return jsonify({'error': 'Attachment not found'}), 404

        return send_from_directory(
            app.config['UPLOAD_FOLDER'],
            row['filename'],
            download_name=row['original_filename'],
            mimetype=row['mime_type']
        )


def analyze_query(query: str) -> dict:
    """Analyze user query to extract risk IDs, keywords, and intent.

    Returns dict with:
        - risk_ids: List of risk IDs mentioned (e.g., ['R001', 'R002'])
        - keywords: Key terms for searching
        - intent: 'create', 'read', 'update', 'delete', or 'general'
    """
    import re

    query_lower = query.lower()

    # Extract risk IDs (R001, R002, etc.)
    risk_ids = re.findall(r'\b[Rr]\d{3}\b', query)
    risk_ids = [rid.upper() for rid in risk_ids]

    # Determine intent
    intent = 'general'
    if any(word in query_lower for word in ['create', 'add', 'new', 'build', 'make', 'write']):
        intent = 'create'
    elif any(word in query_lower for word in ['delete', 'remove', 'drop']):
        intent = 'delete'
    elif any(word in query_lower for word in ['update', 'change', 'modify', 'edit', 'set']):
        intent = 'update'
    elif any(word in query_lower for word in ['show', 'list', 'get', 'find', 'what', 'tell me', 'read']):
        intent = 'read'

    # Extract keywords (filter out common words)
    stop_words = {'the', 'a', 'an', 'is', 'are', 'was', 'were', 'be', 'been', 'being',
                  'have', 'has', 'had', 'do', 'does', 'did', 'will', 'would', 'could',
                  'should', 'may', 'might', 'must', 'shall', 'can', 'to', 'of', 'in',
                  'for', 'on', 'with', 'at', 'by', 'from', 'as', 'into', 'through',
                  'and', 'or', 'but', 'if', 'then', 'else', 'when', 'where', 'why',
                  'how', 'all', 'each', 'every', 'both', 'few', 'more', 'most', 'other',
                  'some', 'such', 'no', 'nor', 'not', 'only', 'own', 'same', 'so',
                  'than', 'too', 'very', 'just', 'about', 'this', 'that', 'these',
                  'those', 'am', 'it', 'its', 'me', 'my', 'i', 'you', 'your', 'we',
                  'our', 'they', 'their', 'them', 'what', 'which', 'who', 'whom',
                  'create', 'add', 'new', 'build', 'make', 'show', 'list', 'get', 'find',
                  'tell', 'please', 'help', 'want', 'need', 'like', 'using', 'use'}

    words = re.findall(r'\b[a-zA-Z]{3,}\b', query_lower)
    keywords = [w for w in words if w not in stop_words]

    return {
        'risk_ids': risk_ids,
        'keywords': keywords,
        'intent': intent
    }


def build_smart_context(query: str, context: dict) -> str:
    """Build a smart, query-relevant system prompt using RAG approach.

    1. Always search library for relevant guidance
    2. Search RACM for relevant risks based on query
    3. Build focused prompt with just relevant data
    """
    analysis = analyze_query(query)

    # 1. Search audit library for relevant guidance
    library_results = []
    if analysis['keywords']:
        search_query = ' '.join(analysis['keywords'][:5])  # Use top 5 keywords
        library_results = db.search_library_keyword(search_query, limit=3)

    # 2. Get relevant RACM data
    relevant_risks = []

    # First, get any specifically mentioned risks
    for risk_id in analysis['risk_ids']:
        risk = db.get_risk(risk_id)
        if risk:
            relevant_risks.append(risk)

    # Then search for keyword-relevant risks (if not too many already)
    if len(relevant_risks) < 5 and analysis['keywords']:
        all_risks = context.get('risks', [])
        for risk in all_risks:
            if len(relevant_risks) >= 5:
                break
            if risk in relevant_risks:
                continue
            # Check if any keyword matches risk description or control
            risk_text = f"{risk.get('risk', '')} {risk.get('control_id', '')}".lower()
            if any(kw in risk_text for kw in analysis['keywords']):
                relevant_risks.append(risk)

    # 3. Build the smart prompt
    audit_summary = context.get('audit_summary', {})
    prompt = f"""You are Felix, an intelligent AI audit assistant with FULL ACCESS to the SmartPapers audit database.

## Current Audit Summary:
- **Annual Audit Plan:** {audit_summary.get('total', 0)} audits planned - By Quarter: {json.dumps(audit_summary.get('by_quarter', {}))}
- **Audit Status:** {json.dumps(audit_summary.get('by_status', {}))}
- **Risks:** {context['risk_summary']['total']} total - {json.dumps(context['risk_summary']['by_status'])}
- **Tasks:** {context['task_summary']['total']} total
- **Issues:** {context['issue_summary']['total']} total
- **Flowcharts:** {context['flowchart_count']} | **Test Documents:** {context['test_doc_count']}

## Annual Audit Plan (All Planned Audits):
"""
    # Add all audits from the plan
    audits = context.get('audits', [])
    if audits:
        for audit in audits:
            prompt += f"- **{audit.get('title', 'Untitled')}** [{audit.get('quarter', 'N/A')}] - {audit.get('status', 'unknown')} | Owner: {audit.get('owner', 'Unassigned')} | Area: {audit.get('audit_area', 'N/A')}\n"
    else:
        prompt += "No audits in the annual plan yet.\n"

    # Add relevant library guidance
    if library_results:
        prompt += "\n## Relevant Guidance from Audit Library:\n"
        for result in library_results:
            doc_name = result.get('document_name', 'Unknown')
            section = result.get('section', '')
            content = result.get('content', '')[:500]  # Limit content length
            prompt += f"\n**{doc_name}**"
            if section:
                prompt += f" - {section}"
            prompt += f"\n{content}...\n"

    # Add relevant risks
    if relevant_risks:
        prompt += "\n## Relevant Risks/Controls:\n"
        for risk in relevant_risks:
            prompt += f"\n**{risk.get('risk_id', 'N/A')}** [{risk.get('status', 'Unknown')}]\n"
            prompt += f"- Risk: {risk.get('risk', 'No description')[:200]}\n"
            prompt += f"- Control: {risk.get('control_id', 'No control')[:200]}\n"
    elif context.get('risks'):
        # If no relevant risks found, show brief overview
        prompt += "\n## RACM Overview (use execute_sql for details):\n"
        for r in context['risks'][:10]:
            prompt += f"- {r.get('risk_id', 'N/A')}: {r.get('risk', '')[:50]}... [{r.get('status', '')}]\n"
        if len(context['risks']) > 10:
            prompt += f"... and {len(context['risks']) - 10} more\n"

    # Add capabilities
    prompt += """
## Your Capabilities - USE TOOLS TO TAKE ACTION:
**CREATE:** Add risks, tasks, issues, flowcharts, test documents
**READ:** Query data with SQL, read documents, search audit library
**UPDATE:** Change statuses, update documentation
**DELETE:** Remove risks, tasks, issues (confirm first)

## Control Writing Standards:
When creating or updating controls, they MUST be concise and follow this format:
- **WHO** does **WHAT**, **WHY**, **WHERE**, and **WHEN**
- Maximum 5 sentences
- No bullet points or detailed procedures in the control description
- Example: "The IT Security Manager reviews user access quarterly using the IAM system to ensure appropriate access levels are maintained and terminated employees are removed within 24 hours of departure."

## Guidelines:
1. **Take Action:** Use tools to do what the user asks, don't just describe
2. **Use Library Guidance:** The audit library content above is from your reference documents - use it to inform your responses
3. **Keep Controls Concise:** Controls must be brief (5 sentences max) - save detail for test procedures
4. **Avoid Duplicates:** Check relevant risks above before creating new ones
5. **Ask if Unclear:** Use ask_clarifying_question when you need more details

Be concise and professional. Use markdown formatting."""

    return prompt


def call_felix_ai(messages, attachments=None):
    """Call Claude API for Felix chat with full tool support."""
    if not ANTHROPIC_API_KEY:
        raise Exception('ANTHROPIC_API_KEY not configured')

    client = anthropic.Anthropic(api_key=ANTHROPIC_API_KEY)

    # Get the latest user message for smart context
    latest_query = ""
    for msg in reversed(messages):
        if msg.get('role') == 'user':
            content = msg.get('content', '')
            if isinstance(content, str):
                latest_query = content
            elif isinstance(content, list):
                # Handle list of content blocks
                for block in content:
                    if isinstance(block, dict) and block.get('type') == 'text':
                        latest_query = block.get('text', '')
                        break
            break

    # Build smart context based on query
    context = db.get_full_context()
    system_prompt = build_smart_context(latest_query, context) if latest_query else build_ai_system_prompt(context)

    # Add attachments context if any
    if attachments:
        attachments_context = "\n\n## Uploaded Documents (this conversation)\nThe user has uploaded the following documents:\n"
        for att in attachments:
            text = att.get('extracted_text', '').strip()
            if text:
                if len(text) > 10000:
                    text = text[:10000] + "\n... [truncated]"
                attachments_context += f"\n### {att['original_filename']}\n```\n{text}\n```\n"
            else:
                attachments_context += f"\n### {att['original_filename']}\n[No text content extracted - may be an image]\n"
        system_prompt += attachments_context

    # Get tools
    tools = get_ai_tools()

    # Call Claude with tools - use higher max_tokens for complex multi-tool requests
    response = client.messages.create(
        model=CLAUDE_MODEL,
        max_tokens=4096,
        system=system_prompt,
        tools=tools,
        messages=messages
    )

    # Process response and execute tools
    final_response = ""
    tool_status_messages = []  # Brief status messages for tools used
    tool_results = []
    current_messages = list(messages)
    max_iterations = 10  # Safety limit to prevent infinite loops

    # Friendly names for tool status messages
    tool_status_names = {
        "search_audit_library": "🔍 Searching audit library...",
        "execute_sql": "🔍 Querying database...",
        "get_audit_summary": "📊 Getting audit summary...",
        "add_racm_row": "➕ Adding risk/control...",
        "update_racm_status": "✏️ Updating status...",
        "create_kanban_task": "📋 Creating task...",
        "create_flowchart": "📐 Creating flowchart...",
        "create_test_document": "📝 Creating test document...",
        "read_test_document": "📖 Reading test document...",
        "read_flowchart": "📐 Reading flowchart...",
        "update_issue_documentation": "💾 Saving documentation...",
        "read_issue_documentation": "📖 Reading documentation...",
    }

    iteration = 0
    while response.stop_reason == "tool_use" and iteration < max_iterations:
        iteration += 1
        for content in response.content:
            if content.type == "tool_use":
                # Add brief status message for this tool
                status = tool_status_names.get(content.name, f"⚙️ Using {content.name}...")
                if status not in tool_status_messages:
                    tool_status_messages.append(status)

                try:
                    result = execute_tool(content.name, content.input)
                    tool_results.append({
                        "type": "tool_result",
                        "tool_use_id": content.id,
                        "content": result
                    })
                except Exception as e:
                    # Log error but continue with other tools
                    error_msg = f"Error executing {content.name}: {str(e)}"
                    logger.error(f"[Felix AI] {error_msg}")
                    tool_results.append({
                        "type": "tool_result",
                        "tool_use_id": content.id,
                        "content": error_msg,
                        "is_error": True
                    })
            # Skip intermediate text blocks - we only want the final response

        # Convert response.content to serializable format for messages
        assistant_content = []
        for content in response.content:
            if content.type == "tool_use":
                assistant_content.append({
                    "type": "tool_use",
                    "id": content.id,
                    "name": content.name,
                    "input": content.input
                })
            elif content.type == "text":
                assistant_content.append({
                    "type": "text",
                    "text": content.text
                })

        current_messages.append({"role": "assistant", "content": assistant_content})
        current_messages.append({"role": "user", "content": tool_results})

        response = client.messages.create(
            model=CLAUDE_MODEL,
            max_tokens=4096,
            system=system_prompt,
            tools=tools,
            messages=current_messages
        )
        tool_results = []

    # Extract final text response
    for content in response.content:
        if hasattr(content, 'text'):
            final_response += content.text

    # Prepend tool status messages if any tools were used
    if tool_status_messages:
        status_block = "\n".join(tool_status_messages) + "\n\n---\n\n"
        final_response = status_block + final_response

    return final_response


# ==================== ADMIN ROUTES ====================

@app.route('/admin')
@require_admin
def admin_dashboard():
    """Admin dashboard page."""
    ctx = get_audit_context()
    return render_template('admin/dashboard.html', active_page='admin', **ctx)


@app.route('/admin/users')
@require_admin
def admin_users():
    """User management - redirects to main admin page."""
    return redirect(url_for('admin_dashboard'))


@app.route('/admin/audits')
@require_admin
def admin_audits():
    """Audit membership management - redirects to main admin page."""
    return redirect(url_for('admin_dashboard'))


# ==================== ADMIN API: Users ====================

@app.route('/api/admin/users', methods=['GET'])
@require_admin
def api_admin_get_users():
    """Get all users."""
    users = db.get_all_users()
    return jsonify(users)


@app.route('/api/admin/users', methods=['POST'])
@require_admin
def api_admin_create_user():
    """Create a new user."""
    data = request.get_json()

    email = data.get('email', '').strip().lower()
    name = data.get('name', '').strip()
    password = data.get('password', '')
    is_admin = int(data.get('is_admin', 0))

    if not email or not name or not password:
        return jsonify({'error': 'Email, name, and password are required'}), 400

    # Check if email already exists
    if db.get_user_by_email(email):
        return jsonify({'error': 'A user with this email already exists'}), 400

    password_hash = hash_password(password)
    user_id = db.create_user(email, name, password_hash, is_active=1, is_admin=is_admin)

    return jsonify({'status': 'created', 'id': user_id})


@app.route('/api/admin/users/<int:user_id>', methods=['GET'])
@require_admin
def api_admin_get_user(user_id):
    """Get a single user."""
    user = db.get_user_by_id(user_id)
    if not user:
        return jsonify({'error': 'User not found'}), 404

    # Remove password hash from response
    return jsonify({
        'id': user['id'],
        'email': user['email'],
        'name': user['name'],
        'is_active': user['is_active'],
        'is_admin': user['is_admin'],
        'created_at': user['created_at'],
        'updated_at': user['updated_at']
    })


@app.route('/api/admin/users/<int:user_id>', methods=['PUT'])
@require_admin
def api_admin_update_user(user_id):
    """Update a user."""
    data = request.get_json()

    updates = {}
    if 'name' in data:
        updates['name'] = data['name'].strip()
    if 'email' in data:
        updates['email'] = data['email'].strip().lower()
    if 'is_active' in data:
        updates['is_active'] = int(data['is_active'])
    if 'is_admin' in data:
        updates['is_admin'] = int(data['is_admin'])
    if 'password' in data and data['password']:
        updates['password_hash'] = hash_password(data['password'])

    if not updates:
        return jsonify({'error': 'No valid updates provided'}), 400

    if db.update_user(user_id, **updates):
        return jsonify({'status': 'updated'})
    else:
        return jsonify({'error': 'User not found'}), 404


@app.route('/api/admin/users/<int:user_id>', methods=['DELETE'])
@require_admin
def api_admin_delete_user(user_id):
    """Deactivate a user (soft delete)."""
    # Don't allow deleting yourself
    current = get_current_user()
    if current and current['id'] == user_id:
        return jsonify({'error': 'Cannot delete your own account'}), 400

    if db.update_user(user_id, is_active=0):
        return jsonify({'status': 'deactivated'})
    else:
        return jsonify({'error': 'User not found'}), 404


@app.route('/api/admin/users/<int:user_id>/memberships', methods=['GET'])
@require_admin
def api_admin_get_user_memberships(user_id):
    """Get all audit memberships for a user."""
    memberships = db.get_user_memberships(user_id)
    return jsonify(memberships)


# ==================== ADMIN API: Roles ====================

@app.route('/api/admin/roles', methods=['GET'])
@require_admin
def api_admin_get_roles():
    """Get all available roles."""
    roles = db.get_all_roles()
    return jsonify(roles)


# ==================== ADMIN API: Audit Memberships ====================

@app.route('/api/admin/audits/<int:audit_id>/memberships', methods=['GET'])
@require_admin
def api_admin_get_audit_memberships(audit_id):
    """Get all users assigned to an audit."""
    memberships = db.get_audit_memberships(audit_id)
    return jsonify(memberships)


@app.route('/api/admin/audits/<int:audit_id>/memberships', methods=['POST'])
@require_admin
def api_admin_add_audit_membership(audit_id):
    """Add a user to an audit."""
    data = request.get_json()

    user_id = data.get('user_id')
    role_id = data.get('role_id')

    if not user_id or not role_id:
        return jsonify({'error': 'user_id and role_id are required'}), 400

    # Check if membership already exists
    if db.get_audit_membership(user_id, audit_id):
        # Update existing membership role
        db.update_audit_membership(audit_id, user_id, role_id)
        return jsonify({'status': 'updated'})

    membership_id = db.add_audit_membership(audit_id, user_id, role_id)
    return jsonify({'status': 'created', 'id': membership_id})


@app.route('/api/admin/audits/<int:audit_id>/memberships/<int:user_id>', methods=['DELETE'])
@require_admin
def api_admin_remove_audit_membership(audit_id, user_id):
    """Remove a user from an audit."""
    if db.remove_audit_membership(audit_id, user_id):
        return jsonify({'status': 'removed'})
    else:
        return jsonify({'error': 'Membership not found'}), 404


# ==================== AUDIT TEAM API (Multi-assignment) ====================

@app.route('/api/admin/audits/<int:audit_id>/team', methods=['GET'])
@require_admin
def api_admin_get_audit_team(audit_id):
    """Get all team members (auditors and reviewers) for an audit."""
    team = db.get_audit_team(audit_id)
    auditors = [m for m in team if m['team_role'] == 'auditor']
    reviewers = [m for m in team if m['team_role'] == 'reviewer']
    viewers = db.get_audit_viewers_list(audit_id) if hasattr(db, 'get_audit_viewers_list') else []
    return jsonify({
        'auditors': auditors,
        'reviewers': reviewers,
        'viewers': viewers
    })


@app.route('/api/admin/audits/<int:audit_id>/team', methods=['POST'])
@require_admin
def api_admin_add_team_member(audit_id):
    """Add a user to an audit team (as auditor or reviewer)."""
    data = request.get_json()
    user = get_current_user()

    user_id = data.get('user_id')
    team_role = data.get('team_role')  # 'auditor' or 'reviewer'

    if not user_id:
        return jsonify({'error': 'user_id is required'}), 400

    if team_role not in ('auditor', 'reviewer'):
        return jsonify({'error': 'team_role must be "auditor" or "reviewer"'}), 400

    result = db.add_to_audit_team(audit_id, user_id, team_role, assigned_by=user['id'])

    if result == -1:
        return jsonify({'status': 'already_exists', 'message': 'User already has this role on this audit'})

    return jsonify({'status': 'created', 'id': result})


@app.route('/api/admin/audits/<int:audit_id>/team/<int:user_id>/<team_role>', methods=['DELETE'])
@require_admin
def api_admin_remove_team_member(audit_id, user_id, team_role):
    """Remove a user from an audit team role."""
    if team_role not in ('auditor', 'reviewer'):
        return jsonify({'error': 'team_role must be "auditor" or "reviewer"'}), 400

    if db.remove_from_audit_team(audit_id, user_id, team_role):
        return jsonify({'status': 'removed'})
    else:
        return jsonify({'error': 'Team membership not found'}), 404


@app.route('/api/admin/audits/<int:audit_id>/viewers', methods=['POST'])
@require_admin
def api_admin_add_viewer(audit_id):
    """Add a viewer to an audit."""
    data = request.get_json()
    user = get_current_user()

    user_id = data.get('user_id')
    if not user_id:
        return jsonify({'error': 'user_id is required'}), 400

    result = db.add_viewer_to_audit(audit_id, user_id, granted_by=user['id'])

    if result == -1:
        return jsonify({'status': 'already_exists', 'message': 'User is already a viewer'})

    return jsonify({'status': 'created', 'id': result})


@app.route('/api/admin/audits/<int:audit_id>/viewers/<int:user_id>', methods=['DELETE'])
@require_admin
def api_admin_remove_viewer(audit_id, user_id):
    """Remove a viewer from an audit."""
    if db.remove_viewer_from_audit(audit_id, user_id):
        return jsonify({'status': 'removed'})
    else:
        return jsonify({'error': 'Viewer not found'}), 404


if __name__ == '__main__':
    app.run(debug=True, port=8001)
